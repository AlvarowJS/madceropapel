"""
@autor Dirección Regional de Transformación Digital
@titularidad GOBIERNO REGIONAL CAJAMARCA
@licencia Python Software Foundation License. Ver LICENCIA.txt
@año 2022
"""

import base64
import json
import os
import shutil
import uuid
from datetime import datetime
from io import BytesIO

import magic
import py7zr
from babel.dates import format_date
from django.conf import settings
from django.db.models import Q, F, Case, When, Value, CharField, BooleanField, OuterRef, Subquery, Count, IntegerField
from django.db.models.functions import Cast, Concat, Coalesce, Substr, StrIndex, Upper
from django.http import HttpResponse
from django.template.defaultfilters import pluralize
from django.template.loader import get_template, render_to_string
from django.urls import reverse
from django.utils import timezone
from django.views.generic import TemplateView, DetailView
from docx import Document
from docx.enum.text import WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Cm
from mailmerge import MailMerge
from rest_framework import views, status
from rest_framework.authentication import TokenAuthentication
from rest_framework.authtoken.models import Token
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from xhtml2pdf import pisa

from apps.inicio.models import Tablero, Anio
from apps.inicio.vistas.inicio import TemplateValidaLogin
from apps.organizacion.models import Area, Dependencia, PeriodoTrabajo, DocumentoTipoArea, TrabajadoresActuales
from apps.sockets.vistas.mensajes import SocketDocGen, SocketMsg
from apps.tramite.formularios.documentodetalle import FormDestino, FormFirma, FormReferencia
from apps.tramite.formularios.documentoemitir import DocumentoEmitirForm, DocumentoEstadoForm, DocumentoAnularForm, \
    DocumentoObservarForm, DocumentoEmitirCambiarResponsableForm
from apps.tramite.formularios.mesapartesregistrar import MesaPartesRegistrarForm
from apps.tramite.models import Documento, Destino, DocumentoPlantilla, DocumentoPDF, DocumentoFirma, \
    DocumentoFirmaEstado, DocumentoEstado, DestinoEstado, DocumentoReferencia, DocumentoReferenciaOrigen, Expediente, \
    AnexoFirma, Anexo, DestinoEstadoMensajeria, MensajeriaModoEntrega
from apps.tramite.vistas.plantillas.proveido import fetch_resources
from modulos.utiles.clases.crud import VistaCreacion, VistaEdicion
from modulos.utiles.clases.varios import Titular


def ObtenerNombreAnio(anio, modo):
    _result = ""
    reg = Anio.objects.filter(numero=anio).first()
    modo = "denominacion" if modo == "N" else "nombredecenio"
    if not reg:
        anio -= 1
        reg = Anio.objects.filter(numero=anio).first()
    if reg:
        _result = eval("reg.%s" % modo)
    return _result


def UnirDocumento(documento, doc, destino, destinofields=False):
    dep = Dependencia.objects.get(codigo=settings.CONFIG_APP["Dependencia"])
    urlverifica = settings.CONFIG_APP["WebVerificador"]

    # Con Copia
    copias = ""
    # copias = "c.c.\n"
    # for copia in doc.des_documento.filter(tipotramite__codigo="1", tipodestinatario="UO"):
    #     if doc.documentotipoarea.area == copia.periodotrabajo.area:
    #         copias += copia.periodotrabajo.iniciales
    #     else:
    #         copias += copia.periodotrabajo.area.nombrecorto
    #     copias += "\n"
    # if doc.origentipo == "O":
    #     copias += doc.documentotipoarea.area.nombrecorto
    # elif doc.origentipo == "P":
    #     copias += doc.responsable.iniciales

    # Datos de la Referencia de Respuesta
    ref_nombre = ""
    ref_asunto = ""
    referencias = doc.referencias.filter(destino__isnull=False)
    if referencias.count() > 0:
        referencia = referencias.first()
        ref_nombre = referencia.descripcion
        ref_asunto = referencia.destino.documento.asunto
    # Días de atención
    atenciondias = ""
    if doc.diasatencion > 0:
        atenciondias = "%s día%s" % (
            doc.diasatencion, pluralize(doc.diasatencion)
        )
    # =====================================
    destino_genero = ""
    destino_nombres = ""
    destino_apellidos = ""
    destino_cargo = ""
    destino_entidad = ""
    destino_entidad_ruc = ""
    destino_ubigeo = ""
    destino_direccion = ""
    destino_ciudad = "Presente"
    if not destinofields:
        destino_genero = destino.periodotrabajo.persona.sexo if destino.tipodestinatario == "UO" else \
            (destino.dependencia_responsable.sexo if destino.tipodestinatario == "DP" else
             (destino.persona.sexo if destino.persona else ""))
        destino_genero = "a" if destino_genero == "F" else ""
        destino_nombres = destino.periodotrabajo.persona.nombres if destino.tipodestinatario == "UO" else \
            (destino.dependencia_responsable.nombres if destino.tipodestinatario == "DP" else
             (destino.persona.nombres if destino.persona else ""))
        destino_apellidos = destino.periodotrabajo.persona.apellidos() if destino.tipodestinatario == "UO" else \
            (destino.dependencia_responsable.apellidos() if destino.tipodestinatario == "DP" else
             (destino.persona.apellidos() if destino.persona else ""))
        destino_cargo = destino.periodotrabajo.Cargo() if destino.tipodestinatario == "UO" else \
            (destino.personacargo if destino.tipodestinatario == "PJ" else
             (destino.DependenciaResponsableCargo() if destino.tipodestinatario == "DP" else destino.personacargo)
             )
        destino_entidad = destino.periodotrabajo.area.nombre if destino.tipodestinatario == "UO" else \
            ((destino.personajuridica.nombrecomercial or destino.personajuridica.razonsocial) if
             destino.tipodestinatario == "PJ" else
             ("%s\n%s" % (
                 destino.dependencia_area_nombre,
                 destino.dependencia.nombre
             ) if destino.tipodestinatario == "DP" else "")
             )
        destino_entidad_ruc = destino.personajuridica.ruc if destino.tipodestinatario == "PJ" else ""
        if destino.ubigeo:
            destino_ubigeo = destino.ubigeo.RutaDepartamento()
            if destino.ubigeo != doc.responsable.area.dependencia.ubigeo:
                destino_ciudad = "%s - %s" % (
                    destino.ubigeo.nombre.upper(),
                    destino.ubigeo.provincia.departamento.nombre.upper()
                )
        if destino.tipodestinatario == "UO" and destino.periodotrabajo.area.rindentepadre:
            destino_entidad = "%s\n%s" % (
                destino_entidad,
                destino.periodotrabajo.area.rindentepadre.nombre
            )
        destino_direccion = destino.direccion or ""
    subtitulo = doc.responsable.area.nombre
    subtitulo2 = " "
    if doc.responsable.area.paracomisiones:
        subtitulo2 = subtitulo
        if doc.responsable.area.padre.paracomisiones:
            subtitulo = doc.responsable.area.dependencia.nombre.upper()
        else:
            subtitulo = doc.responsable.area.padre.nombre
    elif doc.responsable.area.esrindente or doc.responsable.area.rindentepadre:
        if doc.responsable.area.esrindente:
            dep = doc.responsable.area
            subtitulo = doc.responsable.area.nombre
        else:
            dep = doc.responsable.area.rindentepadre
            subtitulo = doc.responsable.area.rindentepadre.nombre
        if doc.responsable.area.rindentepadre:
            subtitulo2 = doc.responsable.area.nombre
    elif doc.responsable.area.nivel > 2:
        if doc.responsable.area.areatipo.codigo == "FU":
            if doc.responsable.area.padre.padre:
                subtitulo = doc.responsable.area.padre.padre.nombre
                subtitulo2 = doc.responsable.area.padre.nombre
            else:
                subtitulo = doc.responsable.area.padre.nombre
        else:
            subtitulo2 = subtitulo
            subtitulo = doc.responsable.area.padre.nombre
    elif doc.responsable.area.areatipo.codigo == "FU":
        if doc.responsable.area.padre.padre:
            if doc.responsable.area.padre.nivel == 2:
                subtitulo = doc.responsable.area.padre.nombre
            else:
                subtitulo = doc.responsable.area.padre.padre.nombre
                subtitulo2 = doc.responsable.area.padre.nombre
        else:
            subtitulo = doc.responsable.area.padre.nombre
    else:
        if doc.responsable.area.padre:
            subtitulo2 = subtitulo
            subtitulo = doc.responsable.area.padre.nombre
    # Nombres del Año
    anio1 = ObtenerNombreAnio(doc.anio, "N").upper()
    anio2 = ObtenerNombreAnio(doc.anio, "D").upper()
    # =====================================
    documento.merge(
        **{
            # Encabezado
            "MAD3_EXPEDIENTE_NUMERO": doc.expedientenro,
            "MAD3_DOCUMENTO_TITULO": doc.responsable.area.dependencia.nombre.upper(),
            "MAD3_DOCUMENTO_SUBTITULO": subtitulo,
            "MAD3_DOCUMENTO_SUBTITULO2": subtitulo2,
            "MAD3_DOCUMENTO_NOMBRE_ANIO2": anio2,
            "MAD3_DOCUMENTO_NOMBRE_ANIO1": anio1,
            # Destinatario
            "MAD3_DESTINATARIO_GENERO": "%s" % destino_genero,
            "MAD3_DESTINATARIO_APELLIDOS": "%s" % destino_apellidos,
            "MAD3_DESTINATARIO_NOMBRES": ", %s" % Titular(destino_nombres),
            "MAD3_DESTINATARIO_CARGO": "%s" % (destino_cargo or ""),
            "MAD3_DESTINATARIO_ENTIDAD":
                "" if len(destino_entidad or "") == 0 else
                "%s%s" % (
                    "\n" if len(destino_cargo or "") > 0 else "",
                    destino_entidad
                ),
            "MAD3_DESTINATARIO_RUC": destino_entidad_ruc,
            "MAD3_DESTINATARIO_UBIGEO": destino_ubigeo,
            "MAD3_DESTINATARIO_DIRECCION":
                "" if len(destino_direccion or "") == 0 else
                "%s%s" % (
                    "\n" if len((destino_cargo or "") + (destino_entidad or "")) > 0 else "",
                    destino_direccion
                ),
            "MAD3_DESTINATARIO_CIUDAD": "%s" % destino_ciudad,
            # Cuerpo
            "MAD3_DOCUMENTO_ASUNTO": "%s" % doc.asunto,
            "MAD3_DOCUMENTO_FECHA": format_date(
                doc.fecha, format="dd 'de' MMMM 'de' yyyy", locale="es"
            ),
            "MAD3_DOCUMENTO_NUMERO": doc.obtenerNumeroSiglas(),
            "MAD3_DOCUMENTO_ATENCION_DIAS": atenciondias,
            # Referencias
            "MAD3_REFERENCIA_NOMBRE": "%s" % ref_nombre,
            "MAD3_REFERENCIA_ASUNTO": "%s" % ref_asunto,
            # EMISOR
            "MAD3_EMITE_APELLIDOS": "%s" % (
                doc.responsable.area.jefeactual.persona.apellidos()
                if doc.responsable.tipo in ["EN", "EP"] and doc.origentipo == "O"
                else doc.responsable.persona.apellidos()
            ),
            "MAD3_EMITE_NOMBRES": ", %s" % Titular(
                doc.responsable.area.jefeactual.persona.nombres
                if doc.responsable.tipo in ["EN", "EP"] and doc.origentipo == "O"
                else doc.responsable.persona.nombres
            ),
            "MAD3_EMITE_NOMBRES_2": "%s" % str(
                doc.responsable.area.jefeactual.persona.nombres
                if doc.responsable.tipo in ["EN", "EP"] and doc.origentipo == "O"
                else doc.responsable.persona.nombres
            ).upper(),
            "MAD3_EMITE_CARGO": doc.responsable.area.jefeactual.Cargo()
            if doc.responsable.tipo in ["EN", "EP"] and doc.origentipo == "O"
            else (doc.responsable.CargoPeriodo() if doc.origentipo == "P" else doc.responsable.Cargo()),
            "MAD3_EMITE_DEPENDENCIA": "%s" % doc.responsable.area.nombre,
            "MAD3_EMITE_INICIALES": "%s" % doc.responsable.iniciales,
            # Pie de Página
            "MAD3_DOCUMENTO_COPIA": copias,
            "MAD3_DEPENDENCIA_DIRECCION": "%s" % dep.direccion,
            "MAD3_DEPENDENCIA_TELEFONO": "%s" % dep.telefono,
            "MAD3_DEPENDENCIA_WEBSITE_TEXTO": "%s" % dep.web.split("//")[1],
            "MAD3_DEPENDENCIA_WEBSITE_URL": "%s" % dep.web.split("//")[1],
            "MAD3_DOCUMENTO_URL_VERIFICA": urlverifica,
            "MAD3_DOCUMENTO_URL_CODIGO": doc.clave,
            "MAD3_SISTEMA_NOMBRE": "%s v.%s" % (settings.CONFIG_APP["TituloCorto"], settings.CONFIG_APP["Version"])
        }
    )
    documento.merge_rows(
        "MAD3_DESTINO_NOMBRE",
        list(doc.des_documento.exclude(ultimoestado__estado="AN").filter(tipodestinatario="UO").order_by("pk").annotate(
            MAD3_DESTINO_NOMBRE=Concat(
                F("periodotrabajo__area__nombre"),
                Value("\n"),
                F("periodotrabajo__persona__apellidocompleto")
            ),
            MAD3_DESTINO_TRAMITE=F("proveido__nombre"),
            MAD3_DESTINO_INDICACION=F("indicacion")
        ).values(
            "MAD3_DESTINO_NOMBRE", "MAD3_DESTINO_TRAMITE", "MAD3_DESTINO_INDICACION"
        ))
    )
    return documento


def PonerReferencias(document, doc):
    ponerReferencia = False
    for pg in document.paragraphs:
        if pg.text.lower() == "[referencia]":
            ponerReferencia = True
            break
    if ponerReferencia:
        pgAsunto = None
        for pg in document.paragraphs:
            if pg.text.lower().startswith("asunto"):
                if doc.referencias.count() > 0:
                    run = pg.add_run()
                    run.add_break(WD_BREAK.LINE)
            elif pg.text.lower() == "[referencia]":
                if doc.referencias.count() > 0:
                    #
                    reftabs = json.loads(
                        doc.documentotipoarea.documentotipo.documentotipoplantilla_set.first(
                        ).referenciatabs
                    )
                    #
                    pg.text = ""
                    run = pg.add_run()
                    run.bold = False
                    run.font.name = "Arial Narrow"
                    run.font.size = Pt(10)
                    if reftabs.get("left_indent"):
                        pg.paragraph_format.left_indent = Cm(reftabs["left_indent"])
                    if reftabs.get("first_line_indent"):
                        pg.paragraph_format.first_line_indent = Cm(reftabs["first_line_indent"])
                    tabs = pg.paragraph_format.tab_stops
                    tabs.clear_all()
                    if reftabs.get("tab_stop"):
                        tabs.add_tab_stop(
                            Cm(reftabs["tab_stop"]), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES
                        )
                    run.add_text("Referencia")
                    run.add_tab()
                    run.add_text(":")
                    run.add_tab()
                    for referencia in doc.referencias.order_by("pk"):
                        if referencia.expedientenro == "0" or referencia.origen.codigo == "EXT":
                            run.add_text("%s" % referencia.descripcion)
                        else:
                            run.add_text("%s (%s: %s)" % (
                                referencia.descripcion,
                                referencia.origen.siglas,
                                referencia.expedientenro
                            ))
                        run.add_break(WD_BREAK.LINE)
                else:
                    pg.text = ""
                break


def ListaDes(doc):
    listdes = list(
        doc.des_documento.exclude(ultimoestado__estado="AN").order_by("pk").annotate(
            des_codigo=Cast("pk", output_field=CharField()),
            des_forma=F("tipotramite__nombre"),
            des_formacodigo=F("tipotramite__codigo"),
            des_nombres=Case(
                When(
                    tipodestinatario="UO",
                    then=F("periodotrabajo__persona__nombres")
                ),
                When(
                    tipodestinatario="PJ",
                    then=Case(
                        When(
                            persona__isnull=True,
                            then=Case(
                                When(
                                    personajuridica__nombrecomercial__isnull=True,
                                    then=Case(
                                        When(
                                            personajuridica__nombrecomercial__isnull=True,
                                            then=F("personajuridica__razonsocial")
                                        ),
                                        default=F("personajuridica__nombrecomercial")
                                    )
                                ),
                                default=F("personajuridica__nombrecomercial")
                            )
                        ),
                        default=F("persona__nombres")
                    )
                ),
                When(tipodestinatario="CI", then=F("persona__nombres")),
                default=Value(""), output_field=CharField()
            ),
            des_apellidos=Case(
                When(
                    tipodestinatario="UO",
                    then=Concat(
                        F("periodotrabajo__persona__paterno"),
                        Value(" "),
                        F("periodotrabajo__persona__materno")
                    )
                ),
                When(
                    tipodestinatario="PJ",
                    then=Case(
                        When(
                            persona__isnull=False,
                            then=Concat(
                                F("persona__paterno"),
                                Value(" "),
                                F("persona__materno")
                            )
                        ),
                        default=Value("")
                    )
                ),
                When(
                    tipodestinatario="CI",
                    then=Concat(
                        F("persona__paterno"),
                        Value(" "),
                        F("persona__materno")
                    )
                ),
                default=Value(""), output_field=CharField()
            ),
            des_cargo=Case(
                When(
                    tipodestinatario="UO",
                    then=Case(
                        When(
                            Q(periodotrabajo__tipo__in=["EN", "EP"]) | Q(periodotrabajo__esjefemodo__in=["TE", "TI"]),
                            then=Concat(
                                Case(
                                    When(
                                        periodotrabajo__persona__sexo="F",
                                        then=F("periodotrabajo__area__cargooficial__nombref")
                                    ),
                                    default=F("periodotrabajo__area__cargooficial__nombrem")
                                ),
                                Case(
                                    When(
                                        Q(periodotrabajo__poscargo__isnull=True)
                                        |
                                        Q(periodotrabajo__esjefemodo="TE"),
                                        then=Value("")
                                    ),
                                    default=Concat(
                                        Value(" "),
                                        F("periodotrabajo__poscargo")
                                    )
                                ),
                                Case(
                                    When(periodotrabajo__esjefemodo="TE", then=Value(" (e)")),
                                    default=Value("")
                                )
                            )
                        ),
                        default=Concat(
                            Case(
                                When(periodotrabajo__persona__sexo="F", then=F("periodotrabajo__cargo__nombref")),
                                default=F("periodotrabajo__cargo__nombrem")
                            ),
                            Case(
                                When(
                                    Q(periodotrabajo__poscargo__isnull=True) | Q(periodotrabajo__esjefemodo="TE"),
                                    then=Value("")
                                ),
                                default=Concat(
                                    Value(" "),
                                    F("periodotrabajo__poscargo")
                                )
                            ),
                            Case(
                                When(periodotrabajo__esjefemodo="TE", then=Value(" (e)")),
                                default=Value("")
                            )
                        )
                    )
                ),
                When(
                    tipodestinatario="PJ",
                    then=Coalesce(
                        F("personacargo"),
                        Case(
                            When(
                                personajuridica__tipo="R",
                                then=Concat(
                                    Value("RUC: "),
                                    F("personajuridica__ruc")
                                )
                            ),
                            default=Value("")
                        )
                    )
                ),
                When(
                    tipodestinatario="CI",
                    then=Concat(
                        F("persona__tipodocumentoidentidad__abreviatura"),
                        Value(": "),
                        F("persona__numero")
                    )
                ),
                default=Value(""), output_field=CharField()
            ),
            des_lugar=Case(
                When(
                    tipodestinatario="UO",
                    then=F("periodotrabajo__area__nombre")
                ),
                When(
                    tipodestinatario="PJ",
                    then=Case(
                        When(
                            persona__isnull=False,
                            then=Case(
                                When(
                                    personajuridica__nombrecomercial__isnull=True,
                                    then=F("personajuridica__razonsocial")
                                ),
                                default=F("personajuridica__nombrecomercial")
                            )
                        ),
                        default=Value("Persona Jurídica")
                    )
                ),
                When(
                    tipodestinatario="CI",
                    then=Value("Ciudadano")
                ),
                default=Value(""), output_field=CharField()
            ),
            des_direccion=Case(
                When(
                    tipodestinatario="UO",
                    then=Value("")
                ),
                When(tipodestinatario="PJ", then=Coalesce(F("direccion"), F("personajuridica__direccion"), Value(""))),
                When(tipodestinatario="CI", then=Coalesce(F("direccion"), F("persona__direccion"), Value(""))),
                default=Value(""), output_field=CharField()
            ),
            des_ciudad=Case(
                When(
                    ubigeo=doc.responsable.area.dependencia.ubigeo,
                    then=Value("Presente")
                ),
                When(
                    ubigeo__isnull=False,
                    then=Concat(
                        Upper(F("ubigeo__nombre")),
                        Value(" - "),
                        Upper(F("ubigeo__provincia__departamento__nombre"))
                    )
                ),
                default=Value("Presente")
            )
        ).values(
            "des_codigo", "des_forma", "des_formacodigo",
            "des_nombres", "des_apellidos", "des_cargo", "des_lugar",
            "des_direccion", "des_ciudad"
        )
    )
    for des in listdes:
        if des["des_apellidos"]:
            des["des_nombres"] = Titular(des["des_nombres"])
    return listdes


def ListaDesMultiple(doc):
    lista = list(
        doc.des_documento.exclude(ultimoestado__estado="AN").filter(
            tipotramite__codigo="0"
        ).order_by("pk").annotate(
            des_codigo=Cast("pk", output_field=CharField()),
            des_nombres=Case(
                When(tipodestinatario="UO", then=F("periodotrabajo__persona__nombres")),
                When(
                    tipodestinatario="PJ", then=Case(
                        When(
                            persona__isnull=True,
                            then=Case(
                                When(
                                    personajuridica__nombrecomercial__isnull=True,
                                    then=F("personajuridica__razonsocial")
                                ),
                                default=F("personajuridica__nombrecomercial")
                            )
                        ),
                        default=F("persona__nombres")
                    )
                ),
                When(tipodestinatario="CI", then=F("persona__nombres")),
                default=Value(""), output_field=CharField()
            ),
            des_apellidos=Case(
                When(
                    tipodestinatario="UO",
                    then=Concat(
                        F("periodotrabajo__persona__paterno"),
                        Value(" "),
                        F("periodotrabajo__persona__materno"),
                    )
                ),
                When(
                    tipodestinatario="PJ", then=Case(
                        When(
                            persona__isnull=True,
                            then=Case(
                                When(
                                    personajuridica__nombrecomercial__isnull=True,
                                    then=F("personajuridica__razonsocial")
                                ),
                                default=F("personajuridica__nombrecomercial")
                            )
                        ),
                        default=Concat(
                            F("persona__paterno"),
                            Value(" "),
                            F("persona__materno"),
                        )
                    )
                ),
                When(
                    tipodestinatario="CI",
                    then=Concat(
                        F("persona__paterno"),
                        Value(" "),
                        F("persona__materno"),
                    )
                ),
                default=Value(""), output_field=CharField()
            ),
            des_cargo=Case(
                When(
                    tipodestinatario="UO",
                    then=Case(
                        When(
                            Q(periodotrabajo__tipo__in=["EN", "EP"]) | Q(periodotrabajo__esjefemodo__in=["TE", "TI"]),
                            then=Concat(
                                Case(
                                    When(
                                        periodotrabajo__persona__sexo="F",
                                        then=F("periodotrabajo__area__cargooficial__nombref")
                                    ),
                                    default=F("periodotrabajo__area__cargooficial__nombrem")
                                ),
                                Case(
                                    When(
                                        Q(periodotrabajo__poscargo__isnull=True)
                                        |
                                        Q(periodotrabajo__esjefemodo="TE"),
                                        then=Value("")
                                    ),
                                    default=Concat(
                                        Value(" "),
                                        F("periodotrabajo__poscargo")
                                    )
                                ),
                                Case(
                                    When(periodotrabajo__esjefemodo="TE", then=Value(" (e)")),
                                    default=Value("")
                                )
                            )
                        ),
                        default=Concat(
                            Case(
                                When(periodotrabajo__persona__sexo="F", then=F("periodotrabajo__cargo__nombref")),
                                default=F("periodotrabajo__cargo__nombrem")
                            ),
                            Case(
                                When(
                                    Q(periodotrabajo__poscargo__isnull=True) | Q(periodotrabajo__esjefemodo="TE"),
                                    then=Value("")
                                ),
                                default=Concat(
                                    Value(" "),
                                    F("periodotrabajo__poscargo")
                                )
                            ),
                            Case(
                                When(periodotrabajo__esjefemodo="TE", then=Value(" (e)")),
                                default=Value("")
                            )
                        )
                    )
                ),
                When(tipodestinatario="PJ", then=Coalesce(
                    F("personacargo"),
                    Case(
                        When(
                            personajuridica__tipo="R",
                            then=Concat(Value("RUC: "), F("personajuridica__ruc"))
                        ),
                        default=Value("")
                    )
                )),
                When(
                    tipodestinatario="CI", then=Concat(
                        F("persona__tipodocumentoidentidad__abreviatura"),
                        Value(": "),
                        F("persona__numero")
                    )
                ),
                default=Value(""), output_field=CharField()
            ),
            des_lugar=Case(
                When(tipodestinatario="UO", then=F("periodotrabajo__area__nombre")),
                When(
                    tipodestinatario="PJ",
                    then=Case(
                        When(
                            persona__isnull=False,
                            then=Case(
                                When(
                                    personajuridica__nombrecomercial__isnull=True,
                                    then=F("personajuridica__razonsocial")
                                ),
                                default=F("personajuridica__nombrecomercial")
                            )
                        ),
                        default=Value("Persona Jurídica")
                    )
                ),
                When(tipodestinatario="CI", then=Value("Ciudadano")),
                default=Value(""), output_field=CharField()
            ),
        ).values(
            "des_codigo", "des_nombres", "des_apellidos", "des_cargo", "des_lugar"
        )
    )
    for registro in lista:
        registro["des_nombres"] = Titular(registro["des_nombres"])
    return lista


def EmitirDocumento(doc, request):
    _result = False
    if not doc.AnexosNoFirmados() > 0:
        _result = True
        docest = DocumentoEstado(
            documento=doc,
            estado="EM",
            creador=request.user
        )
        docest.save()
        # Referencias de Respuesta
        for referencia in doc.referencias.filter(destino__isnull=False):
            crearestado = False
            if not referencia.destino.ultimoestado:
                crearestado = True
            elif referencia.destino.ultimoestado.estado != "AT":
                crearestado = True
            if crearestado:
                DestinoEstado.objects.create(
                    destino=referencia.destino,
                    estado="AT",
                    creador=request.user
                )
        # Notificamos al usuario y al creador
        SocketMsg(
            userid=request.user.pk,
            funcpost='refrescarTableros("dbDespacho,dbEmitidos", true)'
        )
        SocketMsg(
            userid=request.user.pk,
            funcpost='refrescarTabla("refrescar_tabladbRecepcionadosO");'
        )
        SocketMsg(
            userid=request.user.pk,
            funcpost='refrescarTabla("refrescar_tabladbRecepcionadosP");'
        )
        if doc.creador != request.user:
            SocketMsg(
                userid=doc.creador.pk,
                funcpost='refrescarTableros("dbDespacho,dbEmitidos", true)'
            )
            SocketMsg(
                userid=doc.creador.pk,
                funcpost='refrescarTabla("refrescar_tabladbRecepcionadosO");'
            )
            SocketMsg(
                userid=doc.creador.pk,
                funcpost='refrescarTabla("refrescar_tabladbRecepcionadosP");'
            )
        # Notificamos a los destinos de unidades organizacionales
        for destino in doc.des_documento.exclude(ultimoestado__estado="AN"):
            DestinoEstado.objects.create(
                destino=destino,
                estado="NL",
                creador=request.user
            )
            if destino.tipodestinatario == "UO":
                if destino.periodotrabajo.persona.usuario:
                    SocketMsg(
                        tipo='primary',
                        clase='bg-danger',
                        userid=destino.periodotrabajo.persona.usuario.pk,
                        titulo='Bandeja de Entrada',
                        mensaje='%s, %s' % (
                            "Hola %s" % destino.periodotrabajo.persona.nombres,
                            "le han remitido un nuevo documento."
                        ),
                        funcpost='refrescarTableros("dbEntrada", true)'
                    )
            else:
                if destino.mesapartesmodoenvio <= 1:
                    estado = destino.mensajeriamodoentrega.mensajeriaestado
                else:
                    estado = "PE"
                DestinoEstadoMensajeria.objects.create(
                    destino=destino,
                    estado=estado,
                    creador=request.user
                )
    return _result


def GenerarPlantillaAutomatica(doc, context, request):
    doc.save(generardoc=True)
    doc = Documento.objects.get(pk=doc.pk)
    docplla = doc.documentoplantilla
    document = MailMerge(BytesIO(docplla.plantilla))
    destino = docplla.documento.des_documento.exclude(ultimoestado__estado="AN").order_by(
        "tipotramite__codigo", "creado"
    ).first()
    fileTemDoc = os.path.join(settings.TEMP_DIR, uuid.uuid4().hex + ".docx")
    document.write(fileTemDoc)
    doctmp = Document(fileTemDoc)
    os.remove(fileTemDoc)
    docio = BytesIO()
    doctmp.save(docio)
    docplla.contenido = docio.getvalue()
    docplla.save()
    docpdf = docplla.documentopdf_set.first()
    if not docpdf:
        docpdf = DocumentoPDF(
            documentoplantilla=docplla,
            destino=destino,
            pdf=docio.getvalue(),
            creador=request.user
        )
        docpdf.save()
    else:
        docpdf.estado = "G"
    context["documento"] = doc
    template = get_template("tramite/plantillas/%s.html" % doc.documentotipoarea.documentotipo.codigo)
    html = template.render(request=request, context=context)
    result = BytesIO()
    pisa.pisaDocument(html.encode("UTF-8"), dest=result, link_callback=fetch_resources)
    docpdf.pdf = result.getvalue()
    docpdf.save()
    doc.folios = doc.des_documento.first().Folios()
    doc.save()


class DocumentoEmitir(TemplateValidaLogin, VistaCreacion):
    template_name = "tramite/documento/emitir/emitir.html"
    model = Documento
    form_class = DocumentoEmitirForm
    refexp = []
    atenciones = None

    def get_initial(self):
        initial = super(DocumentoEmitir, self).get_initial()
        arearesponsable = self.request.user.persona.periodotrabajoactual(self.request.session.get("cambioperiodo")).area
        if arearesponsable.areatipo.codigo == "FU":
            arearesponsable = arearesponsable.padre
        initial["arearesponsable"] = arearesponsable
        atender = self.request.session.get("atender")
        if atender:
            forma = atender.get("forma")
            dependencia = Dependencia.objects.get(codigo=settings.CONFIG_APP["Dependencia"])
            if atender.get("destinoid"):
                destinos = Destino.objects.filter(pk=atender["destinoid"])
            else:
                if forma == "A":
                    destinos = Destino.objects.filter(pk__in=eval(atender["destinos"]))
                else:
                    destinos = Destino.objects.none()
            if destinos.count() > 0:
                origen = DocumentoReferenciaOrigen.objects.get(codigo="MCP")
                referencias = []
                self.refexp = []
                for destino in destinos:
                    if destinos.count() == 1:
                        initial["asunto"] = destino.documento.asunto
                    verRef = "verRefPdf(\"%s\", \"%s\", \"%s\", \"%s\")" % (
                        origen.codigo,
                        destino.expedientenro,
                        0,
                        destino.pk
                    )
                    if destino.expedientenro not in self.refexp:
                        self.refexp.append(destino.documento.expediente.pk)
                    referencias.append({
                        "codigo": -int(datetime.now().timestamp()),
                        "origenfull": "%s%s%s" % (
                            origen.siglas,
                            "-" if origen.pidedependencia else "",
                            dependencia.siglas
                            if origen.pidedependencia else ""
                        ),
                        "anio": destino.documento.anio,
                        "aniodesc": "%s" % (
                            destino.documento.anio if origen.pideanio else "-"
                        ),
                        "editar": True if destino.documento.ultimoestado.estado in ["RG", "PY", "NL", "LE"] else False,
                        "descripcion": destino.documento.obtenerNumeroSiglas(),
                        "descripcionlink": "%s%s%s" % (
                            "<a href='javascript:;' onclick='%s'>" % verRef if origen.tienepdf else "",
                            destino.documento.obtenerNumeroSiglas(),
                            "</a>" if origen.tienepdf else ""
                        ),
                        "numero": destino.documento.expediente.numero,
                        "origen": origen.pk,
                        "dependencia": dependencia.pk,
                        "expedientenro": destino.expedientenro,
                        "modoref": 1,
                        "oficina": destino.documento.documentotipoarea.area.pk,
                        "oficinanombre": destino.documento.documentotipoarea.area.nombrecorto,
                        "documentotipo": destino.documento.documentotipoarea.documentotipo.pk,
                        "documentotiponombre": destino.documento.documentotipoarea.documentotipo.nombre,
                        "destino": destino.pk
                    })
                initial["referencias"] = json.dumps(referencias)
            elif forma == "D":
                self.atenciones = atender["destinos"]
                initial["atenciones"] = atender["destinos"]
            del self.request.session["atender"]
        return initial

    def get_form(self, form_class=None):
        form = super(DocumentoEmitir, self).get_form(form_class)
        if len(self.refexp) > 1:
            form.fields["expediente"].queryset = Expediente.objects.filter(pk__in=self.refexp)
        else:
            if form.fields.get("expediente"):
                del form.fields["expediente"]
        periodoactual = self.request.user.persona.periodotrabajoactual(self.request.session.get("cambioperiodo"))
        arearesponsable = Area.objects.all()
        filtro = (
                Q(pk=periodoactual.area.pk)
                |
                Q(
                    pk__in=periodoactual.proyeccionareas.values_list("areaorigen__pk")
                )
        )
        if periodoactual.area.areatipo.codigo == "FU":
            filtro = filtro | Q(pk=periodoactual.area.padre.pk)
            arearesponsable = arearesponsable.exclude(areatipo__codigo="FU")
        form.fields["arearesponsable"].queryset = arearesponsable.filter(filtro).order_by("nombre")
        if self.kwargs.get("tipo") == "P":
            del form.fields["arearesponsable"]
            del form.fields["responsable"]
            form.fields["documentotipoarea"].widget.dependent_fields = {}
            form.fields["documentotipoarea"].queryset = DocumentoTipoArea.objects.filter(
                area=periodoactual.area,
                documentotipo__usoprofesional=True
            ).order_by("documentotipo__nombre")
        else:
            form.fields["documentotipoarea"].queryset = DocumentoTipoArea.objects.order_by("documentotipo__nombre")
        if not self.atenciones:
            del form.fields["atenciones"]
        else:
            form.fields["documentotipoarea"].queryset = form.fields["documentotipoarea"].queryset.filter(
                documentotipo__paraderivacion=True
            )
        return form

    def get_context_data(self, **kwargs):
        context = super(DocumentoEmitir, self).get_context_data(**kwargs)
        if self.request.POST:
            self.template_name = "tramite/documento/emitir/emitir_form.html"
        else:
            tiposdestinos = Destino.TIPODESTINATARIO
            if self.kwargs.get("tipo") in ["P"] and not self.request.user.persona.periodotrabajoactual(
                    self.request.session.get("cambioperiodo")).emiteexterno:
                tiposdestinos = [tiposdestinos[0]]
            context["tiposdestinos"] = tiposdestinos
        context["ModosEntregas"] = MensajeriaModoEntrega.objects.filter(estado=True).order_by("orden")
        context["se_puede_editar"] = True
        return context

    def form_valid(self, form):
        if form.is_valid():
            form.instance.origentipo = self.kwargs.get("tipo")
            periodoactual = self.request.user.persona.periodotrabajoactual(
                self.request.session.get("cambioperiodo")
            )
            if self.kwargs.get("tipo") == "P":
                upt = self.request.user.persona.ultimoperiodotrabajo
                if periodoactual.area == upt.area:
                    periodoactual = upt
                form.instance.responsable = periodoactual
                form.instance.emisor = periodoactual
                form.instance.documentotipoarea = DocumentoTipoArea.objects.filter(
                    documentotipo=form.cleaned_data["documentotipoarea"].documentotipo,
                    area=periodoactual.area
                ).first()
            else:
                form.instance.documentotipoarea = DocumentoTipoArea.objects.filter(
                    documentotipo=form.cleaned_data["documentotipoarea"].documentotipo,
                    area=form.cleaned_data["arearesponsable"]
                ).first()
                form.instance.emisor = periodoactual
            atenciones = self.request.POST.get("atenciones")
            if atenciones:
                form.instance.expediente = Destino.objects.get(pk=eval(atenciones)[0]).documento.expediente
            super(DocumentoEmitir, self).form_valid(form)
            context = self.get_context_data(form=form)
            doc = form.instance
            if not settings.CONFIG_APP.get("EnDesarrollo"):
                doc.confidencial = False
                doc.save()
            # 01 - Guardamos los destinos
            idsdestinos = []
            cldt = form.cleaned_data
            destinos = json.loads(cldt["destinos"])
            for destino in destinos:
                if not destino.get("tieneentregafisica"):
                    destino["entregafisica"] = None
                if destino.get("tipodestinatario") == "PJ" and destino.get("personajuridicatipo") == "O":
                    destino["personajuridica"] = destino.get("personajuridicarz")
                else:
                    destino["personajuridicarz"] = None
                fdest = FormDestino(data=destino, request=self.request)
                if not destino.get("mensajeriamodoentrega"):
                    fdest.instance.mensajeriamodoentrega = MensajeriaModoEntrega.objects.filter(
                        estado=True
                    ).order_by("orden").first()
                if fdest.is_valid():
                    fdest.instance.documento = doc
                    fdest.instance.creador = self.request.user
                    fdest.save()
                    idsdestinos.append({
                        "idold": int(destino["codigo"]),
                        "idnew": fdest.instance.pk
                    })
                else:
                    print(fdest.errors)
            context["idsdestinos"] = idsdestinos
            # 03 - Guardamos las referencias
            idsreferencias = []
            if cldt.get("referencias"):
                referencias = json.loads(cldt["referencias"])
                for referencia in referencias:
                    fref = FormReferencia(data=referencia)
                    if referencia["origenfull"] == "EXT":
                        fref.fields["numero"].required = False
                        fref.fields["expedientenro"].required = False
                        referencia["descripcion"] = referencia["refereciaexterna"]
                    if fref.is_valid():
                        fref.instance.documento = doc
                        fref.instance.creador = self.request.user
                        fref.instance.estado = DocumentoFirmaEstado.objects.get(codigo="SF")
                        fref.instance.motivorechazado = None
                        fref.save()
                        idsreferencias.append({
                            "idold": int(referencia["codigo"]),
                            "idnew": fref.instance.pk
                        })
                        if fref.instance.destino and not doc.expediente:
                            if self.request.POST.get("expediente"):
                                doc.expediente = Expediente.objects.get(pk=self.request.POST.get("expediente"))
                            else:
                                doc.expediente = fref.instance.destino.documento.expediente
                            doc.save()
                    else:
                        print(fref.instance.pk, fref.errors)
            context["idsreferencias"] = idsreferencias
            # 04 - Guardamos las firmas
            idsfirmas = []
            if cldt.get("firmas"):
                firmas = json.loads(cldt["firmas"])
                for firma in firmas:
                    ffirm = FormFirma(data=firma)
                    if ffirm.is_valid():
                        ffirm.instance.documento = doc
                        ffirm.instance.creador = self.request.user
                        ffirm.instance.estado = DocumentoFirmaEstado.objects.get(codigo="SF")
                        ffirm.instance.motivorechazado = None
                        ffirm.save()
                        idsfirmas.append({
                            "idold": int(firma["codigo"]),
                            "idnew": ffirm.instance.pk
                        })
                    else:
                        print(ffirm.instance.pk, ffirm.errors)
            context["idsfirmas"] = idsfirmas
            context["tab"] = Tablero.objects.filter(codigo="dbEnProyecto").first()
            context["tabid"] = doc.id
            # 00 - Verificamos si el tipo de documento necesita número
            if doc.documentotipoarea.documentotipo.plantillaautomatica:
                doc = Documento.objects.get(pk=doc.pk)
                GenerarPlantillaAutomatica(doc, context, self.request)
            if atenciones:
                doc2 = Documento.objects.get(pk=doc.pk)
                docids = [str(doc2.pk)]
                for idx, referencia in enumerate(Destino.objects.filter(pk__in=eval(atenciones))):
                    if idx > 0:
                        doc2.pk = None
                        doc2.ultimoestado = None
                        doc2.numero = None
                        doc2.siglas = None
                        doc2.expediente = referencia.documento.expediente
                        doc2.save()
                    # Le agregamos la referencia
                    fref = FormReferencia(data={
                        'codigo': -1,
                        'origenfull': 'MAD3',
                        'anio': referencia.documento.anio,
                        'aniodesc': str(referencia.documento.anio),
                        'descripcion': referencia.documentonrosiglas,
                        'numero': referencia.documento.expediente.numero,
                        'origen': 3,
                        'dependencia': referencia.documento.expediente.dependencia.id,
                        'expedientenro': referencia.expedientenro,
                        'modoref': 1,
                        'oficina': referencia.documento.documentotipoarea.area.id,
                        'oficinanombre': referencia.documento.documentotipoarea.area.nombrecorto,
                        'documentotipo': referencia.documento.documentotipoarea.documentotipo.id,
                        'documentotiponombre': referencia.documento.documentotipoarea.documentotipo.nombre,
                        'destino': referencia.id
                    })
                    fref.instance.documento = doc2
                    fref.instance.creador = self.request.user
                    fref.instance.estado = DocumentoFirmaEstado.objects.get(codigo="SF")
                    fref.instance.motivorechazado = None
                    fref.save()
                    if idx > 0:
                        # Le copiamos los demás datos:
                        docids.append(str(doc2.id))
                        # Destinos y Estados
                        for destinodoc in doc.des_documento.order_by("pk"):
                            destinodoc2 = Destino.objects.get(pk=destinodoc.id)
                            destinodoc2.id = None
                            destinodoc2.ultimoestado = None
                            destinodoc2.documento = doc2
                            destinodoc2.save()
                        # Firmas
                        for firmadoc in doc.firmas.order_by("pk"):
                            firmadoc2 = DocumentoFirma.objects.get(pk=firmadoc.pk)
                            firmadoc2.id = None
                            firmadoc2.documento = doc2
                            firmadoc2.save()
                        # Plantillas
                    GenerarPlantillaAutomatica(doc2, context, self.request)
                    DocumentoEstado.objects.create(
                        documento=doc2,
                        estado="PD",
                        creador=self.request.user
                    )
                context["docids"] = "_".join(docids)
                SocketMsg(
                    userid=self.request.user.pk,
                    funcpost='refrescarTableros("dbEnDespacho,dbRecepcionados")'
                )
            else:
                SocketMsg(
                    userid=self.request.user.pk,
                    funcpost='refrescarTableros("dbEnProyecto,dbRecepcionados")'
                )
            result = self.render_to_response(context)
        else:
            result = self.render_to_response(self.get_context_data(form=form))
        return result

    def form_invalid(self, form):
        fi = super(DocumentoEmitir, self).form_invalid(form)
        return fi


def VerificaEdicion(documento, request):
    periodoactual = request.user.persona.periodotrabajoactual(request.session.get("cambioperiodo"))
    return documento.ultimoestado.estado == "PY" and (
            documento.creador == request.user or
            documento.responsable == periodoactual or
            (periodoactual.tipo in ["EN", "EP", "AP"] and documento.responsable.area == periodoactual.area) or
            (documento.responsable.tipo in ["EN", "EP"] and periodoactual.esjefe)
    )


class DocumentoEmitirBotones(TemplateValidaLogin, DetailView):
    template_name = "tramite/documento/emitir/botones/_botones.html"
    http_method_names = "post"
    model = Documento

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        context = self.get_context_data(object=self.object)
        context["formid"] = request.POST.get("formid")
        context["tab"] = request.POST.get("tab")
        context["tabid"] = request.POST.get("tabid")
        context["cf"] = False if request.POST.get("cf", "false") == "false" else True
        periodoactual = request.user.persona.periodotrabajoactual(request.session.get("cambioperiodo"))
        context["se_puede_editar"] = VerificaEdicion(self.object, request)
        return self.render_to_response(context)


def get_lista_destinos(documento):
    return list(
        documento.des_documento.exclude(ultimoestado__estado="AN").order_by("pk").annotate(
            codigo=F("pk"),
            cargonombre=Case(
                When(
                    tipodestinatario="UO",
                    then=Case(
                        When(
                            Q(periodotrabajo__tipo__in=["EN", "EP"]) | Q(periodotrabajo__esjefemodo__in=["TE", "TI"]),
                            then=Concat(
                                Case(
                                    When(
                                        periodotrabajo__persona__sexo="F",
                                        then=F("periodotrabajo__area__cargooficial__nombref")
                                    ),
                                    default=F("periodotrabajo__area__cargooficial__nombrem")
                                ),
                                Case(
                                    When(
                                        Q(periodotrabajo__poscargo__isnull=True)
                                        |
                                        Q(periodotrabajo__esjefemodo="TE"),
                                        then=Value("")
                                    ),
                                    default=Concat(
                                        Value(" "),
                                        F("periodotrabajo__poscargo")
                                    )
                                ),
                                Case(
                                    When(periodotrabajo__esjefemodo="TE", then=Value(" (e)")),
                                    default=Value("")
                                )
                            )
                        ),
                        default=Concat(
                            Case(
                                When(periodotrabajo__persona__sexo="F", then=F("periodotrabajo__cargo__nombref")),
                                default=F("periodotrabajo__cargo__nombrem")
                            ),
                            Case(
                                When(
                                    Q(periodotrabajo__poscargo__isnull=True) | Q(periodotrabajo__esjefemodo="TE"),
                                    then=Value("")
                                ),
                                default=Concat(
                                    Value(" "),
                                    F("periodotrabajo__poscargo")
                                )
                            ),
                            Case(
                                When(periodotrabajo__esjefemodo="TE", then=Value(" (e)")),
                                default=Value("")
                            )
                        )
                    )
                ),
                When(
                    tipodestinatario="PJ",
                    then=F("personacargo")
                ),
                When(
                    tipodestinatario="CI",
                    then=Value("Ciudadano")
                ),
                default=Value(""),
                output_field=CharField()
            ),
            dependenciasiglas=Value("", output_field=CharField()),
            personajuridicatipo=F("personajuridica__tipo"),
            personajuridicaruc=Case(When(personajuridica__tipo="R", then=F("personajuridica__ruc")), default=Value("")),
            personajuridicarz=Case(
                When(personajuridica__tipo="O", then=F("personajuridica_id")), default=Value(0)
            ),
            area=Cast("periodotrabajo__area_id", output_field=CharField()),
            personadni=F("persona__numero"),
            nombre=Case(
                When(
                    tipodestinatario="UO",
                    then=Concat(
                        Case(
                            When(
                                periodotrabajo__area__rindentepadre__isnull=False,
                                then=F("periodotrabajo__area__rindentepadre__siglas")
                            ),
                            default=F("periodotrabajo__area__dependencia__siglas")
                        ),
                        Value(" - "),
                        F("periodotrabajo__area__nombre"),
                        output_field=CharField()
                    )
                ),
                When(
                    tipodestinatario="PJ",
                    then=Case(
                        When(
                            personajuridica__tipo="R",
                            then=Concat(
                                F("personajuridica__ruc"),
                                Value(" - "),
                                Case(
                                    When(
                                        personajuridica__nombrecomercial__isnull=True,
                                        then=F("personajuridica__razonsocial")
                                    ),
                                    default=F("personajuridica__nombrecomercial")
                                )
                            )
                        ),
                        default=Concat(
                            Value("OTRO - "),
                            F("personajuridica__razonsocial")
                        )
                    )
                ),
                When(
                    tipodestinatario="CI",
                    then=F("persona__numero")
                ),
                When(
                    tipodestinatario="DP",
                    then=Concat(
                        F("dependencia__nombre"),
                        Value(" - "),
                        F("dependencia_area_nombre")
                    )
                ),
                default=Value(""),
                output_field=CharField()
            ),
            dirigidoa=Case(
                When(
                    tipodestinatario="UO",
                    then=F("periodotrabajo__persona__apellidocompleto"),
                ),
                When(
                    tipodestinatario__in=["PJ", "CI"],
                    then=Concat(
                        F("persona__numero"),
                        Value(" - "),
                        F("persona__nombrecompleto")
                    ),
                ),
                When(
                    tipodestinatario__in=["DP"],
                    then=F("dependencia_responsable__apellidocompleto"),
                ),
                default=Value(""),
                output_field=CharField()
            ),
            cargo=Case(
                When(
                    tipodestinatario="UO",
                    then=Case(
                        When(
                            Q(periodotrabajo__tipo__in=["EN", "EP"]) | Q(periodotrabajo__esjefemodo__in=["TE", "TI"]),
                            then=Concat(
                                Case(
                                    When(
                                        periodotrabajo__persona__sexo="F",
                                        then=F("periodotrabajo__area__cargooficial__nombref")
                                    ),
                                    default=F("periodotrabajo__area__cargooficial__nombrem")
                                ),
                                Case(
                                    When(
                                        Q(periodotrabajo__poscargo__isnull=True)
                                        |
                                        Q(periodotrabajo__esjefemodo="TE"),
                                        then=Value("")
                                    ),
                                    default=Concat(
                                        Value(" "),
                                        F("periodotrabajo__poscargo")
                                    )
                                ),
                                Case(
                                    When(periodotrabajo__esjefemodo="TE", then=Value(" (e)")),
                                    default=Value("")
                                )
                            )
                        ),
                        default=Concat(
                            Case(
                                When(periodotrabajo__persona__sexo="F", then=F("periodotrabajo__cargo__nombref")),
                                default=F("periodotrabajo__cargo__nombrem")
                            ),
                            Case(
                                When(
                                    Q(periodotrabajo__poscargo__isnull=True) | Q(periodotrabajo__esjefemodo="TE"),
                                    then=Value("")
                                ),
                                default=Concat(
                                    Value(" "),
                                    F("periodotrabajo__poscargo")
                                )
                            ),
                            Case(
                                When(periodotrabajo__esjefemodo="TE", then=Value(" (e)")),
                                default=Value("")
                            )
                        )
                    )
                ),
                When(
                    tipodestinatario="PJ",
                    then=Case(
                        When(personacargo__isnull=True, then=Value("-")),
                        default=F("personacargo")
                    )
                ),
                When(
                    tipodestinatario="CI",
                    then=Value("Ciudadano")
                ),
                When(
                    tipodestinatario="DP",
                    then=Case(
                        When(dependencia_responsable__sexo="F", then=F("dependencia_responsable_cargo__nombref")),
                        default=F("dependencia_responsable_cargo__nombrem")
                    )
                ),
                default=Value(""),
                output_field=CharField()
            ),
            direccionfull=Case(
                When(
                    tipodestinatario="UO",
                    then=Value("-")
                ),
                When(
                    tipodestinatario__in=["PJ", "CI"],
                    then=Concat(
                        F("direccion"),
                        Value(" / "),
                        F("ubigeo__provincia__departamento__nombre"),
                        Value(" / "),
                        F("ubigeo__provincia__nombre"),
                        Value(" / "),
                        F("ubigeo__nombre"),
                    )
                ),
                When(
                    tipodestinatario__in=["DP"],
                    then=F("dependencia__direccion")
                ),
                default=Value(""),
                output_field=CharField()
            ),
            tramite=F("tipotramite__nombre"),
            proveidotexto=F("proveido__nombre"),
            acciones=Value(""),
            ultest=F("ultimoestado__estado"),
            obsdes=Case(
                When(
                    ultimoestado__estado="RH",
                    then=Concat(
                        Value("<span class='text-danger'>"),
                        F("ultimoestado__observacion"),
                        Value("</span>"),
                        output_field=CharField()
                    )
                ),
                default=Value("")
            ),
            obsnew=Value(""),
            seharechazado=Subquery(DestinoEstado.objects.filter(
                estado="RH", destino__pk=OuterRef("pk")
            ).order_by("-creado").values("pk")[:1]),
            editar=Case(
                When(
                    Q(seharechazado__isnull=True, ultimoestado__estado__in=["RG", "PY", "NL", "LE"]),
                    then=Value(True)
                ),
                default=Value(False),
                output_field=BooleanField()
            ),
            eliminar=Case(
                When(ultimoestado__estado__in=["RG", "PY", "NL", "LE", "RH"], then=Value(True)),
                default=Value(False),
                output_field=BooleanField()
            ),
            modo=Value("destinos"),
            dependencia_responsable_dni=F("dependencia_responsable__numero"),
            dependencia_responsable_texto=Case(
                When(
                    tipodestinatario__in=["DP"],
                    then=Concat(
                        F("dependencia_responsable_dni"),
                        Value(" - "),
                        F("dirigidoa"),
                        Value(" - "),
                        F("cargo")
                    )
                ),
                default=Value("")
            ),
            tieneentregafisica=Case(
                When(Q(entregafisica__isnull=True) | Q(entregafisica=""), then=Value(False)),
                default=Value(True),
                output_field=BooleanField()
            ),
            modoentrega=Case(
                When(
                    tipodestinatario__in=["PJ", "CI"],
                    then=Concat(
                        Value("<dropdown data-color='"),
                        F("mensajeriamodoentrega__color"),
                        Value("' data-titulo='"),
                        F("mensajeriamodoentrega__nombre"),
                        Value("' data-icono='"),
                        F("mensajeriamodoentrega__icono"),
                        Value("'>")
                    )
                ),
                default=Value("")
            )
        ).values(
            "codigo",
            "tipodestinatario",
            "cargonombre",
            "dependenciasiglas",
            "personajuridicatipo",
            "personajuridicaruc",
            "personajuridica",
            "personajuridicarz",
            "ubigeo",
            "direccion",
            "referencia",
            "correo",
            "personadni",
            "persona",
            "personacargo",
            "tipotramite",
            "proveido",
            "obsdes",
            "obsnew",
            "ultest",
            "indicacion",
            "nombre",
            "dirigidoa",
            "cargo",
            "direccionfull",
            "proveidotexto",
            "tramite",
            "editar",
            "eliminar",
            "modo",
            "acciones",
            "area",
            "periodotrabajo",
            "dependencia",
            "dependencia_area",
            "dependencia_area_nombre",
            "dependencia_responsable",
            "dependencia_responsable_texto",
            "dependencia_responsable_dni",
            "dependencia_responsable_cargo_id",
            "diasatencion",
            "entregafisica",
            "tieneentregafisica",
            "mensajeriamodoentrega",
            "modoentrega"
        )
    )


def get_lista_referencias(documento):
    urlanx = reverse("apptra:documento_anexos", kwargs={"pk": 0})[:-1]
    return list(
        documento.referencias.order_by("pk").annotate(
            codigo=F("pk"),
            origenfull=Concat(
                F("origen__siglas"),
                Case(
                    When(origen__pidedependencia=True, then=Value("-")),
                    default=Value("")
                ),
                Case(
                    When(origen__pidedependencia=True, then=F("dependencia__siglas")),
                    default=Value("")
                )
            ),
            aniodesc=Case(
                When(origen__pideanio=True, then=Cast("anio", CharField())),
                default=Value("-"),
                output_field=CharField()
            ),
            editar=Case(
                When(destino__isnull=False, then=Value(False)),
                default=Case(
                    When(documento__ultimoestado__estado__in=["RG", "PY", "NL", "LE"], then=Value(True)),
                    default=Value(False),
                    output_field=BooleanField()
                )
            ),
            eliminar=Case(
                When(destino__isnull=False, then=Value(False)),
                default=Case(
                    When(documento__ultimoestado__estado__in=["RG", "PY", "NL", "LE"], then=Value(True)),
                    default=Value(False),
                    output_field=BooleanField()
                )
            ),
            descripcionlink=Concat(
                Case(
                    When(
                        origen__tienepdf=True,
                        then=Concat(
                            Value("<a href='javascript:;' onclick='"),
                            Value("verRefPdf(\""),
                            F("origen__codigo"),
                            Value("\", \""),
                            F("expedientenro"),
                            Value("\", \""),
                            Case(
                                When(
                                    Q(expedienteemi__isnull=True)
                                    |
                                    Q(expedienteemi=""),
                                    then=Value("0")
                                ),
                                default=F("expedienteemi")
                            ),
                            Value("\", "),
                            Case(
                                When(destino__isnull=True, then=Value("0")),
                                default=Cast("destino_id", CharField()),
                            ),
                            Value(")"),
                            Value("'>")
                        )
                    ),
                    default=Value("")
                ),
                F("descripcion"),
                Case(
                    When(origen__tienepdf=True, then=Value("</a>")),
                    default=Value("")
                ),
                Case(
                    When(
                        Q(
                            origen__codigo="MCP",
                            destino__documento__anexos__isnull=False
                        ),
                        then=Concat(
                            Value("<a class='btn btn-xs btn-icon btn-primary ml-2' href='"),
                            Value(urlanx),
                            Cast("destino__documento_id", CharField()),
                            Value("' title='Anexos' rel='tooltip' "),
                            Value("data-toggle='modal' data-target='#modal-principal-centro' data-modal-size='xl'"),
                            Value("><i class='fa font-size-sm fas fa-paperclip'></i></a>")
                        )
                    ),
                    default=Value(""),
                    output_field=CharField()
                )

            ),
            modoref=F("modo"),
            refereciaexterna=Case(
                When(origen__codigo="EXT", then=F("descripcion")),
                default=Value("")
            ),
            expid=Case(
                When(origen__codigo="MCP", then=F("destino__documento__expediente_id")),
                default=Value("0"), output_field=IntegerField()
            )
        ).distinct().values(
            "codigo",
            "origenfull",
            "anio",
            "aniodesc",
            "descripcion",
            "descripcionlink",
            "numero",
            "origen",
            "dependencia",
            "expedientenro",
            "modo",
            "oficina",
            "oficinanombre",
            "documentotipo",
            "documentotiponombre",
            "destino",
            "editar",
            "eliminar",
            "modoref",
            "refereciaexterna",
            "expid"
        )
    )


def get_lista_firmas(documento):
    whens = [
        When(modo=k, then=Value(v))
        for k, v in dict(DocumentoFirma.MODO).items()

    ]
    return list(
        documento.firmas.order_by("pk").annotate(
            codigo=F("pk"),
            area=F("empleado__area_id"),
            areanombre=F("empleado__area__nombre"),
            empleadonombre=F("empleado__persona__apellidocompleto"),
            modonombre=Case(*whens, default=Value(''), output_field=CharField()),
            modofirma=F("modo"),
            estadonombre=Concat(
                Value("<div class='text-center'>"),
                Value("<i class='"),
                F("estado__icono"),
                Value("' style='color: "),
                F("estado__color"),
                Value("'"),
                Case(
                    When(
                        estado__codigo="RE",
                        then=Concat(
                            Value(" data-toggle='popover' title='Observado' data-html='true' data-content='"),
                            F("motivorechazado"),
                            Value("'")
                        )
                    ),
                    default=Value(""),
                    output_field=CharField()
                ),
                Value("></i></div>"),
                output_field=CharField()
            ),
            editar=Case(
                When(documento__ultimoestado__estado__in=["RG", "PY", "NL", "LE"], then=Value(True)),
                default=Value(False),
                output_field=BooleanField()
            ),
            eliminar=Case(
                When(documento__ultimoestado__estado__in=["RG", "PY", "NL", "LE"], then=Value(True)),
                default=Value(False),
                output_field=BooleanField()
            )
        ).values(
            "codigo",
            "area",
            "areanombre",
            "empleado",
            "empleadonombre",
            "modo",
            "modonombre",
            "modofirma",
            "estadonombre",
            "editar",
            "eliminar"
        )
    )


def get_lista_firmantes(pk, request):
    whens = [
        When(estado=k, then=Value(v))
        for k, v in dict(AnexoFirma.ESTADO).items()
    ]
    return AnexoFirma.objects.filter(
        anexo_id=pk
    ).order_by(
        "pk"
    ).annotate(
        estadonombre=Case(*whens, default=Value(''), output_field=CharField()),
        estadoicono=Case(
            When(estado="FI", then=Value("far fa-check-circle")),
            When(estado="RE", then=Value("far fa-times-circle")),
            default=Value("far fa-clock")
        ),
        estadocolor=Case(
            When(estado="FI", then=Value("success")),
            When(estado="RE", then=Value("danger")),
            default=Value("warning")
        ),
        areasiglas=F("empleado__area__nombrecorto"),
        posnombres=StrIndex("empleado__persona__nombres", Value(" ")),
        personasiglas=Concat(
            Case(
                When(posnombres__gt=0, then=Substr("empleado__persona__nombres", 1, F("posnombres") - 1)),
                default=F("empleado__persona__nombres")
            ),
            Value(" "),
            F("empleado__persona__paterno")
        ),
        personafull=F("empleado__persona__nombrecompleto"),
        areafull=F("empleado__area__nombre"),
        firmar=Case(
            When(
                Q(
                    estado="SF",
                    empleado__pk=request.user.persona.periodotrabajoactual(request.session.get("cambioperiodo")).pk,
                    anexo__documento__ultimoestado__estado="PD"
                ),
                then=Value(True)
            ),
            default=Value(False)
        )
        # ).values(
        #     "id",
        #     "estado",
        #     "estadonombre",
        #     "estadoicono",
        #     "estadocolor",
        #     "areafull",
        #     "areasiglas",
        #     "empleado",
        #     "personafull",
        #     "personasiglas",
        #     "modo",
        #     "firmar"
    )


def get_lista_anexos(documento, request, editar=True):
    lista = list(
        documento.anexos.order_by("pk").annotate(
            editarv=Value(editar),
            codigo=F("pk"),
            firmantes=Value(""),
            acciones=Value(""),
            editar=Case(
                When(editarv=False, then=Value(False)),
                When(documento__ultimoestado__estado__in=["RG", "PY", "NL", "LE"], then=Value(True)),
                default=Value(False),
                output_field=BooleanField()
            ),
            eliminar=Case(
                When(editarv=False, then=Value(False)),
                When(documento__ultimoestado__estado__in=["RG", "PY", "NL", "LE"], then=Value(True)),
                default=Value(False),
                output_field=BooleanField()
            )
        ).values(
            "codigo",
            "descripcion",
            "archivonombre",
            "firmantes",
            "acciones",
            "editar",
            "eliminar"
        )
    )
    for registro in lista:
        registro["firmantes"] = render_to_string(
            "tramite/documento/emitir/emitir_anexo_firmas.html",
            context={
                "firmas": get_lista_firmantes(registro["codigo"], request)
            },
            request=request
        )
        archivo = Anexo.objects.get(pk=registro["codigo"]).archivo
        registro["archivonombre"] = render_to_string(
            "tramite/documento/emitir/emitir_anexo_abrirdoc.html",
            context={
                "anx": Anexo.objects.get(pk=registro["codigo"]),
                "cod": registro["codigo"],
                # "espdf": magic.from_buffer(archivo, mime=True) == "application/pdf",
                "nombre": registro["archivonombre"]
            },
            request=request
        )
    return lista


class DocumentoEmitirEditar(TemplateValidaLogin, VistaEdicion):
    template_name = "tramite/documento/emitir/emitir.html"
    model = Documento
    form_class = DocumentoEmitirForm
    refexp = []

    def setup(self, request, *args, **kwargs):
        result = super(DocumentoEmitirEditar, self).setup(request, *args, **kwargs)
        doc = self.get_object()
        if not doc.origentipo in ["O", "P"]:
            self.form_class = MesaPartesRegistrarForm
        return result

    def get_initial(self):
        initial = super(DocumentoEmitirEditar, self).get_initial()
        if self.object.responsable:
            arearesponsable = self.object.responsable.area
            if arearesponsable.areatipo.codigo == "FU":
                arearesponsable = arearesponsable.padre
            initial["arearesponsable"] = arearesponsable
        destinos = get_lista_destinos(self.object)
        # Referencias
        referencias = get_lista_referencias(self.object)
        self.refexp = []
        for ref in referencias:
            if ref["destino"] and ref["origenfull"] == "MAD3":
                self.refexp.append(ref["expid"])
        # Firmas
        firmas = get_lista_firmas(self.object)
        # Anexos
        anexos = get_lista_anexos(self.object, self.request)
        #
        initial["destinos"] = json.dumps(destinos, ensure_ascii=False)
        initial["referencias"] = json.dumps(referencias, ensure_ascii=False)
        initial["firmas"] = json.dumps(firmas, ensure_ascii=False)
        initial["anexos"] = json.dumps(anexos, ensure_ascii=False)
        return initial

    def get_form(self, form_class=None):
        form = super(DocumentoEmitirEditar, self).get_form(form_class)
        periodoactual = self.request.user.persona.periodotrabajoactual(self.request.session.get("cambioperiodo"))
        if len(self.refexp) > 1:
            form.fields["expediente"].queryset = Expediente.objects.filter(pk__in=self.refexp)
        else:
            if form.fields.get("expediente"):
                del form.fields["expediente"]
        if "arearesponsable" in form.fields:
            if self.object.numero:
                del form.fields["arearesponsable"]
                del form.fields["documentotipoarea"]
                if not self.get_object().documentotipoarea.documentotipo.tieneforma:
                    del form.fields["forma"]
                form.fields["responsable"].widget.dependent_fields = {}
                form.fields["responsable"].queryset = form.fields["responsable"].queryset.filter(
                    area=self.object.responsable.area
                )
            else:
                arearesponsable = Area.objects.all()
                filtro = (
                        Q(pk=periodoactual.area.pk)
                        |
                        Q(pk__in=periodoactual.proyeccionareas.values_list("areaorigen__pk"))
                )
                if periodoactual.area.areatipo.codigo == "FU":
                    filtro = filtro | Q(pk=periodoactual.area.padre.pk)
                    arearesponsable = arearesponsable.exclude(areatipo__codigo="FU")
                form.fields["arearesponsable"].queryset = arearesponsable.filter(filtro).order_by("nombre")
        if self.object.origentipo == "P":
            if form.fields.get("arearesponsable"):
                del form.fields["arearesponsable"]
            del form.fields["responsable"]
            if form.fields.get("documentotipoarea"):
                form.fields["documentotipoarea"].widget.dependent_fields = {}
                form.fields["documentotipoarea"].queryset = DocumentoTipoArea.objects.filter(
                    area=periodoactual.area,
                    documentotipo__usoprofesional=True
                )
        else:
            if form.fields.get("documentotipoarea"):
                form.fields["documentotipoarea"].queryset = DocumentoTipoArea.objects.order_by("documentotipo__nombre")
        if form.fields.get("atenciones"):
            del form.fields["atenciones"]
        return form

    def get_context_data(self, **kwargs):
        context = super(DocumentoEmitirEditar, self).get_context_data(**kwargs)
        periodoactual = self.request.user.persona.periodotrabajoactual(self.request.session.get("cambioperiodo"))
        if self.request.POST:
            self.template_name = "tramite/documento/emitir/emitir_form.html"
        else:
            tiposdestinos = Destino.TIPODESTINATARIO
            if self.get_object().origentipo in ["P"] and not periodoactual.emiteexterno:
                tiposdestinos = [tiposdestinos[0]]
            context["tiposdestinos"] = tiposdestinos
        tab = self.kwargs.get("tab", "")
        tabid = self.kwargs.get("tabid", None)
        context["tab"] = tab
        context["tabid"] = tabid
        if tab.lower() in ["dbentrada", "dbrecepcionados", "dbrechazados"]:
            context["destino"] = Destino.objects.filter(pk=tabid).first()
        elif tab.lower() in ["dbfirmavb", "dbdespacho"]:
            context["firmasvbs"] = self.object.anexos.filter(
                firmadores__empleado=periodoactual,
                firmadores__estado="SF"
            ).count()
        context["se_puede_editar"] = VerificaEdicion(self.object, self.request)
        context["MESAPARTESMODOENVIO"] = Destino.MESAPARTESMODOENVIO[1:]
        context["mnanexos"] = [
            {
                "nombre": "Simple",
                "url": reverse("apptra:documento_anexo_agregar", kwargs={"pk": self.object.pk})
            },
            {
                "nombre": "Múltiple",
                "url": reverse("apptra:documento_anexo_masivo_agregar", kwargs={"pk": self.object.pk})
            },
        ]
        if self.request.session.get("docdp_error"):
            context["docdp_error"] = self.request.session["docdp_error"]
            del self.request.session["docdp_error"]
        context["ModosEntregas"] = MensajeriaModoEntrega.objects.filter(estado=True).order_by("orden")
        return context

    def form_valid(self, form):
        if form.is_valid():
            # Cambio de firmador
            # cambiofirma = (self.get_object().responsable != form.cleaned_data["responsable"])
            cambiofirma = False
            if form.instance.origentipo == "P":
                periodoactual = self.request.user.persona.periodotrabajoactual(
                    self.request.session.get("cambioperiodo")
                )
                if form.cleaned_data.get("documentotipoarea"):
                    form.instance.documentotipoarea = DocumentoTipoArea.objects.filter(
                        documentotipo=form.cleaned_data["documentotipoarea"].documentotipo,
                        area=periodoactual.area
                    ).first()
            else:
                if form.cleaned_data.get("documentotipoarea"):
                    form.instance.documentotipoarea = DocumentoTipoArea.objects.filter(
                        documentotipo=form.cleaned_data["documentotipoarea"].documentotipo,
                        area=form.cleaned_data["arearesponsable"]
                    ).first()
            #
            super(DocumentoEmitirEditar, self).form_valid(form)
            doc = self.get_object()
            if not settings.CONFIG_APP.get("EnDesarrollo"):
                doc.confidencial = False
                doc.save()
            cldt = form.cleaned_data
            # ===============================================
            if not hasattr(doc, "documentoplantilla"):
                DocumentoPlantilla.objects.create(
                    documento=doc,
                    plantilla=doc.documentotipoarea.documentotipo.documentotipoplantilla_set.first().archivo,
                    creador=self.request.user
                )
            docplla = doc.documentoplantilla
            if docplla:
                modosave = self.request.POST["modosave"]
                if modosave == "update":
                    # Copia nuevamente la plantilla del Tipo de Documento
                    plantilla = doc.documentotipoarea.documentotipo.documentotipoplantilla_set.first()
                    if plantilla:
                        docplla.plantilla = plantilla.archivo
                if modosave in ["update", "clean"]:
                    cambiofirma = True
                    # Limpia el Contenido para nuevamente hacer Merge
                    docplla.contenido = None
                    docplla.save()
                    if not docplla.documento.documentotipoarea.documentotipo.esmultiple or (
                            docplla.documento.documentotipoarea.documentotipo.esmultiple and
                            docplla.documento.forma == "L"
                    ):
                        try:
                            doc.documentoplantilla.documentopdf_set.all().delete()
                        except Exception as e:
                            pass
                # # -----------
            # =====================================================================
            # 01 - Guardamos los destinos
            destinos = json.loads(cldt["destinos"])
            idsEditados = []
            idsdestinos = []
            for destino in destinos:
                if not destino.get("tieneentregafisica"):
                    destino["entregafisica"] = None
                if destino.get("tipodestinatario") == "PJ" and destino.get("personajuridicatipo") == "O":
                    destino["personajuridica"] = destino.get("personajuridicarz")
                else:
                    destino["personajuridicarz"] = None
                codigo = int(destino["codigo"])
                instance = None
                if codigo > 0:
                    instance = Destino.objects.get(pk=codigo)
                if destino.get("dependencia_responsable_cargo_id"):
                    destino["dependencia_responsable_cargo"] = destino["dependencia_responsable_cargo_id"]
                fdest = FormDestino(data=destino, instance=instance, request=self.request)
                if not destino.get("mensajeriamodoentrega"):
                    fdest.instance.mensajeriamodoentrega = MensajeriaModoEntrega.objects.filter(
                        estado=True
                    ).order_by("orden").first()
                if fdest.is_valid():
                    esnuevo = False
                    if not fdest.instance.pk:
                        esnuevo = True
                        fdest.instance.documento = doc
                        fdest.instance.creador = self.request.user
                    else:
                        fdest.instance.editor = self.request.user
                    idsEditados.append(fdest.instance.pk)
                    try:
                        fdest.save()
                        if esnuevo:
                            idsdestinos.append({
                                "idold": codigo,
                                "idnew": fdest.instance.pk
                            })
                            idsEditados.append(fdest.instance.pk)
                    except Exception as e:
                        print(e)
                        form.add_error(None, "El destinatario ha cambiado")
                else:
                    form.add_error(None, "El destinatario ha cambiado o hay un error")
                    # print(destino["indicacion"])
                    # print(len(destino["indicacion"]))
                    if fdest.instance.pk:
                        idsEditados.append(fdest.instance.pk)
                    # print(fdest.instance.pk, fdest.errors)
            for desdel in doc.des_documento.exclude(ultimoestado__estado="AN").exclude(pk__in=idsEditados):
                if desdel.HaRechazado():
                    destinoestado = DestinoEstado.objects.create(
                        destino=desdel,
                        estado="AN",
                        creador=self.request.user
                    )
                    desdel.ultimoestado = destinoestado
                    desdel.save()
                    if hasattr(desdel, "documentopdf"):
                        desdel.documentopdf.revisiones.all().delete()
                        desdel.documentopdf.tokens.all().delete()
                        desdel.documentopdf.delete()
                else:
                    desdel.ultimoestado = None
                    desdel.ultimoestadomensajeria = None
                    desdel.save()
                    desdel.destinoestados.all().delete()
                    desdel.estadosmensajeria.all().delete()
                    if hasattr(desdel, "documentopdf"):
                        desdel.documentopdf.revisiones.all().delete()
                        desdel.documentopdf.tokens.all().delete()
                        desdel.documentopdf.delete()
                    desdel.delete()
            context = self.get_context_data(form=form)
            context["cambiofirmador"] = cambiofirma
            context["idsdestinos"] = idsdestinos
            # 03 - Guardamos las referencias
            idsRefsEditadas = []
            idsRefsNuevas = []
            if cldt.get("referencias"):
                referencias = json.loads(cldt["referencias"])
                for referencia in referencias:
                    codigo = int(referencia["codigo"])
                    instance = None
                    if codigo > 0:
                        instance = DocumentoReferencia.objects.get(pk=codigo)
                    fref = FormReferencia(data=referencia, instance=instance)
                    if referencia["origenfull"] == "EXT":
                        fref.fields["numero"].required = False
                        fref.fields["expedientenro"].required = False
                        referencia["descripcion"] = referencia["refereciaexterna"]
                    if fref.is_valid():
                        esnuevo = False
                        if not fref.instance.pk:
                            esnuevo = True
                            fref.instance.documento = doc
                            fref.instance.creador = self.request.user
                        else:
                            fref.instance.editor = self.request.user
                        fref.save()
                        idsRefsEditadas.append(fref.instance.pk)
                        if esnuevo:
                            idsRefsNuevas.append({
                                "idold": codigo,
                                "idnew": fref.instance.pk
                            })
                    else:
                        # Aún existiendo error si ya tiene instancia debe pasar a lista
                        if fref.instance.pk:
                            idsRefsEditadas.append(fref.instance.pk)
                        print(fref.instance.pk, fref.errors)
            doc.referencias.exclude(pk__in=idsRefsEditadas).delete()
            context["idsreferencias"] = idsRefsNuevas
            # 04 - Guardamos las firmas
            idsFirmasEditados = []
            idsFirmasNuevos = []
            if cldt.get("firmas"):
                firmas = json.loads(cldt["firmas"])
                for firma in firmas:
                    codigo = int(firma["codigo"])
                    instance = None
                    if codigo > 0:
                        instance = DocumentoFirma.objects.get(pk=codigo)
                    ffirm = FormFirma(data=firma, instance=instance)
                    if ffirm.is_valid():
                        esnuevo = False
                        if not ffirm.instance.pk:
                            esnuevo = True
                            ffirm.instance.documento = doc
                            ffirm.instance.creador = self.request.user
                        else:
                            ffirm.instance.editor = self.request.user
                        ffirm.instance.estado = DocumentoFirmaEstado.objects.get(codigo="SF")
                        ffirm.instance.motivorechazado = None
                        ffirm.save()
                        idsFirmasEditados.append(ffirm.instance.pk)
                        if esnuevo:
                            idsFirmasNuevos.append({
                                "idold": codigo,
                                "idnew": ffirm.instance.pk
                            })
                    else:
                        print(ffirm.instance.pk, ffirm.errors)
                context["idsfirmas"] = idsFirmasNuevos
            for firmdel in doc.firmas.exclude(pk__in=idsFirmasEditados):
                firmdel.delete()
            if not form.non_field_errors():
                # 00 - Verificamos si el tipo de documento necesita número
                if doc.documentotipoarea.documentotipo.plantillaautomatica:
                    doc = Documento.objects.get(pk=doc.pk)
                    GenerarPlantillaAutomatica(doc, context, self.request)
            # ===================================================
            context["form"] = form
            context["cambiofirmador"] = cambiofirma
            #
            result = self.render_to_response(context)
        else:
            result = self.render_to_response(self.get_context_data(form=form))
        return result

    def form_invalid(self, form):
        fi = super(DocumentoEmitirEditar, self).form_invalid(form)
        return fi


class DocumentoEmitirGenerar(TemplateValidaLogin, TemplateView):
    template_name = "tramite/documento/emitir/emitir_generar_docx.html"
    http_method_names = "post"

    def post(self, request, *args, **kwargs):
        context = self.get_context_data(**kwargs)
        return self.render_to_response(context=context)

    def render_to_response(self, context, **response_kwargs):
        if self.request.POST:
            doc = Documento.objects.filter(pk=self.kwargs.get("pk")).first()
            if doc:
                dochost = "%s://%s" % (
                    self.request.scheme,
                    self.request.get_host()
                )
                docplla = doc.documentoplantilla
                docplla.save(genuuid=True)
                token, created = Token.objects.get_or_create(user=self.request.user)
                if doc.documentotipoarea.documentotipo.esmultiple and not \
                        doc.documentotipoarea.documentotipo.esmultipledestino:
                    esnuevo = True
                    if not doc.documentoplantilla.contenido:
                        docplla = doc.documentoplantilla
                        docplla.contenido = doc.documentotipoarea.documentotipo.documentotipoplantilla_set.first().archivoparacontenido
                        docplla.save()
                    elif doc.documentoplantilla.contenido != doc.documentotipoarea.documentotipo.documentotipoplantilla_set.first().archivoparacontenido:
                        esnuevo = False
                else:
                    esnuevo = False
                docJson = {
                    "nombre": doc.nombreDocumento(),
                    "token": token.key,
                    "esmultiple": doc.documentotipoarea.documentotipo.esmultiple,
                    "forma": doc.forma,
                    "esnuevo": esnuevo,
                    "esmultipledestino": doc.documentotipoarea.documentotipo.esmultipledestino,
                    "esprincipal": doc.responsable.area.dependencia.pk == 1,
                    "rutabajada": "%s%s" % (
                        dochost,
                        reverse("apptra:documento_emitir_generar_bajar", kwargs={"codigo": docplla.codigo})
                    ),
                    "rutasubida": "%s%s" % (
                        dochost,
                        reverse("apptra:documento_emitir_generar_subir", kwargs={"codigo": docplla.codigo})
                    ),
                    "rutaerror": "%s%s" % (
                        dochost,
                        reverse("apptra:documento_emitir_generar_error", kwargs={"codigo": docplla.codigo})
                    ),
                    "demo": "MAD3 - DEMO" if settings.CONFIG_APP.get("Test", False) else ""
                }
                docJson = json.dumps(docJson)
                context["codigo"] = docplla.codigo
                docCode = base64.b64encode(docJson.encode("ascii")).decode('ascii')
                context["urldown"] = "%s%s" % (
                    settings.CONFIG_APP["AppOffice"],
                    docCode
                )
        return super(DocumentoEmitirGenerar, self).render_to_response(context, **response_kwargs)


class DocumentoEmitirGenerarBajar(views.APIView):
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        codigo = self.kwargs.get("codigo")
        doc = Documento.objects.filter(documentoplantilla__codigo=codigo).first()
        if doc:
            docplla = doc.documentoplantilla
            multiple = doc.documentotipoarea.documentotipo.esmultiple
            multipledestino = doc.documentotipoarea.documentotipo.esmultipledestino
            forma = doc.forma
            dtp = doc.documentotipoarea.documentotipo.documentotipoplantilla_set.first()
            nombre = "%s.7z" % uuid.uuid4().hex
            destino = docplla.documento.des_documento.exclude(ultimoestado__estado="AN").order_by("creado").first()
            response = HttpResponse(
                status=status.HTTP_200_OK,
                content_type='application/x-7z-compressed'
            )
            response["Content-Disposition"] = "attachment; filename=%s" % nombre
            if not docplla.contenido:
                # Reemplazamos los Fields
                # if dtp.archivoparacontenido:
                #     document = MailMerge(BytesIO(dtp.archivoparacontenido))
                # else:
                document = MailMerge(BytesIO(docplla.plantilla))
                UnirDocumento(document, doc, destino)
                # Creamos el docx con mergefields
                fileTemDoc = os.path.join(settings.TEMP_DIR, uuid.uuid4().hex + ".docx")
                document.write(fileTemDoc)
                document = Document(fileTemDoc)
                os.remove(fileTemDoc)
                # Ponemos las referencias
                PonerReferencias(document, doc)
                docB = BytesIO()
                document.save(docB)
                docplla.contenido = docB.getvalue()
                docplla.save()
            dirtemp = os.path.join(settings.TEMP_DIR, uuid.uuid4().hex)
            os.mkdir(dirtemp)
            contenido = docplla.contenido
            fileWord = open("%s/contenido.docx" % dirtemp, "wb")
            fileWord.write(contenido)
            fileWord.close()
            # if multiple and not multipledestino:  # and forma == "I":
            # Plantilla
            # document = MailMerge(BytesIO(docplla.plantilla))
            # UnirDocumento(document, doc, destino, True)
            # fileTemDoc = os.path.join(settings.TEMP_DIR, uuid.uuid4().hex + ".docx")
            # document.write(fileTemDoc)
            # document = Document(fileTemDoc)
            # os.remove(fileTemDoc)
            # PonerReferencias(document, doc)
            # fileWord = open("%s/plantilla.docx" % dirtemp, "wb")
            # document.save(fileWord)
            # fileWord.close()
            # Pos Firma
            # document = MailMerge(BytesIO(dtp.archivoposfirma))
            # UnirDocumento(document, doc, destino, True)
            # fileTemDoc = "%s/posfirma.docx" % dirtemp
            # document.write(fileTemDoc)
            # document.close()
            if multipledestino:
                # Destinos en JSON
                fileDest = open("%s/destinos.txt" % dirtemp, "w", encoding='utf8')
                fileDest.write(json.dumps({"destinos": ListaDesMultiple(doc)}, ensure_ascii=False))
                fileDest.close()
            elif multiple:
                # Destinos en JSON
                fileDest = open("%s/destinos.txt" % dirtemp, "w", encoding='utf8')
                if forma == "I":
                    campos = (
                        "des_codigo", "des_apellidos", "des_sep", "des_nombres", "des_cargo",
                        "des_lugar", "des_direccion", "des_ciudad"
                    )
                    for idxc, campo in enumerate(campos):
                        fileDest.write(campo)
                        if idxc < len(campos) - 1:
                            fileDest.write("\t")
                    fileDest.write("\n")
                    listam = ListaDes(doc)
                    for idxd, destino in enumerate(listam):
                        for idxc, campo in enumerate(campos):
                            if campo == "des_sep":
                                separador = ", " if destino["des_apellidos"] else ""
                                fileDest.write(separador)
                            elif campo == "des_apellidos" and not destino["des_apellidos"]:
                                fileDest.write(destino["des_nombres"])
                            elif campo == "des_nombres" and not destino["des_apellidos"]:
                                fileDest.write("")
                            else:
                                fileDest.write(destino[campo])
                            if idxc < len(campos) - 1:
                                fileDest.write("\t")
                        if idxd < len(listam) - 1:
                            fileDest.write("\n")
                else:
                    fileDest.write(json.dumps({"destinos": ListaDes(doc)}, ensure_ascii=False))
                fileDest.close()
            file7z = os.path.join("%s/%s.7z" % (settings.TEMP_DIR, uuid.uuid4().hex))
            zipeado = py7zr.SevenZipFile(file7z, "w")
            zipeado.writeall(dirtemp, "")
            zipeado.close()
            shutil.rmtree(dirtemp)
            response.write(open(file7z, "rb").read())
            os.remove(file7z)
            return response
        return Response([], status=status.HTTP_400_BAD_REQUEST)


class DocumentoEmitirGenerarSubir(views.APIView):
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated]
    http_method_names = "post"

    def post(self, request, *args, **kwargs):
        codigo = self.kwargs.get("codigo")
        docplla = DocumentoPlantilla.objects.filter(codigo=codigo).first()
        if docplla:
            archivos = request.FILES.getlist("file")
            archivozip = archivos[0].read()
            if request.headers.get("SaveAs"):
                docplla.contenido = archivozip
                docplla.save()
            else:
                stream7z = py7zr.SevenZipFile(BytesIO(archivozip), "r")
                _folderTmp = os.path.join(settings.TEMP_DIR, uuid.uuid4().hex)
                stream7z.extractall(path=_folderTmp)
                for _file in stream7z.files:
                    _nombre, _extension = os.path.splitext(_file.filename)
                    _fileload = open(os.path.join(_folderTmp, _file.filename), "rb")
                    _fileread = _fileload.read()
                    _fileload.close()
                    if _extension.lower() == ".pdf":
                        if docplla.documento.documentotipoarea.documentotipo.esmultiple and \
                                not docplla.documento.documentotipoarea.documentotipo.esmultipledestino and \
                                docplla.documento.forma == "I":
                            _iddestino = int(_nombre.replace("plantilla_", ""))
                            destino = Destino.objects.filter(pk=_iddestino).first()
                            if destino:
                                archivoPdf = DocumentoPDF.objects.filter(
                                    documentoplantilla=docplla,
                                    destino=destino
                                ).first()
                                if not archivoPdf:
                                    archivoPdf = DocumentoPDF(
                                        documentoplantilla=docplla,
                                        destino=destino,
                                        creador=request.user
                                    )
                                archivoPdf.pdffirmado = None
                                archivoPdf.pdf = _fileread
                                archivoPdf.estado = "G"
                                archivoPdf.actualizado = timezone.now()
                                archivoPdf.editor = request.user
                                archivoPdf.save()
                        else:
                            destino = docplla.documento.des_documento.exclude(
                                ultimoestado__estado="AN"
                            ).order_by("pk").first()
                            archivoPdf = DocumentoPDF.objects.filter(
                                documentoplantilla=docplla,
                                destino=destino
                            ).first()
                            if not archivoPdf:
                                archivoPdf = DocumentoPDF(
                                    documentoplantilla=docplla,
                                    destino=destino,
                                    creador=request.user
                                )
                            archivoPdf.pdffirmado = None
                            archivoPdf.pdf = _fileread
                            archivoPdf.estado = "G"
                            archivoPdf.actualizado = timezone.now()
                            archivoPdf.editor = request.user
                            archivoPdf.save()
                    else:
                        docplla.contenido = _fileread
                        docplla.save()
                shutil.rmtree(_folderTmp)
                docplla.codigo = None
                docplla.save()
                # Notificamos al cliente
                mensaje = 'El documento se ha generado correctamente'
                if docplla.documento.documentotipoarea.documentotipo.esmultiple and not \
                        docplla.documento.documentotipoarea.documentotipo.esmultipledestino and \
                        docplla.documento.forma == "I":
                    mensaje = 'Los documentos se han generado correctamente'
                SocketDocGen("Ok", request.user.id, mensaje)
        return Response([], status=status.HTTP_200_OK)


class DocumentoEmitirGenerarError(views.APIView):
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated]
    http_method_names = "post"

    def post(self, request, *args, **kwargs):
        codigo = self.kwargs.get("codigo")
        docplla = DocumentoPlantilla.objects.filter(codigo=codigo).first()
        if docplla:
            docplla.codigo = None
            docplla.save()
            # Notificamos al cliente
            SocketDocGen("Error", request.user.id, request.POST.get("mensaje"))
        return Response([], status=status.HTTP_200_OK)


class DocumentoEmitirDespacho(TemplateValidaLogin, VistaEdicion):
    template_name = "tramite/documento/emitir/emitir_estado.html"
    model = Documento
    form_class = DocumentoEstadoForm

    def form_valid(self, form):
        context = self.get_context_data(**self.get_form_kwargs())
        if form.is_valid():
            periodoactual = self.request.user.persona.periodotrabajoactual(self.request.session.get("cambioperiodo"))
            codest = form.cleaned_data["codest"]
            modoest = form.cleaned_data.get("modoest", 0) or 0
            doc = self.object
            seguir = True
            if codest == "PD":
                if doc.documentotipoarea.documentotipo.esmultiple and \
                        not doc.documentotipoarea.documentotipo.esmultipledestino and \
                        doc.forma == "I":
                    if doc.des_documento.exclude(ultimoestado__estado="AN").filter(
                            documentopdf__isnull=True
                    ).count() > 0:
                        self.request.session["docdp_error"] = "<div class='mb-3'>%s</div><div>%s</div>" % (
                            "Se encontró un cambio en los destinatarios.",
                            "Vuelva a Generar su documento"
                        )
                        seguir = False
                elif doc.ultimoestado.estado == "PD":
                    seguir = False
                    context["pdupdate"] = True
                if seguir:
                    context["pdok"] = True
            elif codest == "PY":
                if doc.ultimoestado.estado in ["EM", "RP", "RT", "AT", "AR"]:
                    if (doc.ultimoestado.estado == "EM" and (periodoactual.esjefe or periodoactual.tipo in ["EN"])) or \
                            (doc.ultimoestado.estado == "RT" and doc.RechazadoTotal()):
                        # Está emitido y es el jefe de área o encargado
                        # Está recibido totalmente y está rechazado totalmente
                        pass
                    else:
                        context["error"] = "El documento está en estado %s" % doc.ultimoestado.get_estado_display()
                        seguir = False
                elif doc.documentoplantilla.documentopdf_set.filter(pdffirma__isnull=False).count() > 0 and \
                        doc.responsable != periodoactual and doc.ultimoestado.firmado:
                    context["error"] = "El documento ya fue firmado"
                    seguir = False
                elif doc.ultimoestado.estado == "PY":
                    context["pyupdate"] = True
                    seguir = False
                if not seguir:
                    if context.get("error"):
                        context["error"] += " y no puede ser regresado a proyecto por Ud."
                else:
                    context["pyok"] = True
            if context.get("pyok") or context.get("pyupdate") or context.get("pdupdate"):
                context["urledit"] = reverse("apptra:documento_emitir_editar", kwargs={
                    "pk": doc.id,
                    "tab": "dbEnProyecto",
                    "tabid": doc.id
                })
            if seguir:
                for destino in doc.des_documento.exclude(ultimoestado__estado="AN").filter(
                        tipodestinatario__in=["PJ", "CI"]
                ):
                    destino.mesapartesmodoenvio = modoest
                    destino.save()
                if doc.ultimoestado.estado != codest:
                    docest = DocumentoEstado(
                        documento=doc,
                        estado=codest,
                        creador=self.request.user,
                        observacion=doc.ultimoestado.observacion if doc.ultimoestado.estado == "OF" else None
                    )
                    docest.save()
                    doc.ultimoestado = docest
                    doc.save()
                    for docpdf in doc.documentoplantilla.documentopdf_set.all():
                        docpdf.estado = "G"
                        docpdf.pdffirma = None
                        docpdf.save()
                    if codest == "PD":  # Para Despacho
                        if doc.firmas.count() == 0:
                            if doc.responsable.persona.usuario != self.request.user:
                                mensaje = '%s, %s' % (
                                    "Hola %s" % doc.responsable.persona.nombres,
                                    "le han enviado un documento para Firmar."
                                )
                                SocketMsg(
                                    tipo='primary',
                                    clase='bg-primary',
                                    userid=doc.responsable.persona.usuario.pk,
                                    titulo="Bandeja En Proyecto",
                                    mensaje=mensaje,
                                    funcpost='refrescarTableros("dbEnProyecto,dbDespacho", true)'
                                )
                                SocketMsg(
                                    userid=self.request.user.pk,
                                    funcpost='refrescarTableros("dbEnProyecto,dbDespacho", true)'
                                )
                            else:
                                SocketMsg(
                                    userid=self.request.user.pk,
                                    funcpost='refrescarTableros("dbEnProyecto,dbDespacho", true)'
                                )
                                SocketMsg(
                                    userid=doc.responsable.persona.usuario.pk,
                                    funcpost='refrescarTableros("dbEnProyecto,dbDespacho", true)'
                                )
                            if self.request.user != doc.creador:
                                SocketMsg(
                                    userid=doc.creador.pk,
                                    funcpost='refrescarTableros("dbEnProyecto,dbDespacho", true)'
                                )
                        else:
                            SocketMsg(
                                userid=doc.responsable.persona.usuario.pk,
                                funcpost='refrescarTableros("dbEnProyecto,dbDespacho", true)'
                            )
                            if doc.responsable.persona.usuario != self.request.user:
                                SocketMsg(
                                    userid=self.request.user.pk,
                                    funcpost='refrescarTableros("dbEnProyecto,dbDespacho", true)'
                                )
                            # Como se pasa a Despacho enviamos un mensaje a los que tienen que firmar o poner vb
                            for firmador in doc.firmas.all():
                                mensaje = ""
                                if firmador.empleado.persona.usuario != self.request.user:
                                    mensaje = '%s, %s' % (
                                        "Hola %s" % firmador.empleado.persona.nombres,
                                        "le han enviado un documento para %s." % firmador.get_modo_display()
                                    )
                                SocketMsg(
                                    tipo="warning",
                                    clase="bg-warning",
                                    userid=firmador.empleado.persona.usuario.pk,
                                    titulo="Bandeja Firmar o Visar",
                                    mensaje=mensaje,
                                    funcpost='refrescarTableros("dbEnProyecto,dbDespacho,dbFirmaVB", true)'
                                )
                    elif codest == "PY":
                        SocketMsg(
                            userid=doc.responsable.persona.usuario.pk,
                            funcpost='refrescarTableros("dbEnProyecto,dbEmitidos,dbDespacho", true)'
                        )
                        if doc.responsable.persona.usuario != self.request.user:
                            SocketMsg(
                                userid=self.request.user.pk,
                                funcpost='refrescarTableros("dbEnProyecto,dbEmitidos,dbDespacho", true)'
                            )
                        for destino in doc.des_documento.exclude(ultimoestado__estado="AN"):
                            if destino.ultimoestado.estado != "RG":
                                destinoestado = DestinoEstado.objects.create(
                                    destino=destino,
                                    estado="RG",
                                    creador=self.request.user
                                )
                                destino.ultimoestado = destinoestado
                                destino.save()
                        for firmador in doc.firmas.all():
                            if firmador.estado.codigo != "SF":
                                firmador.estado = DocumentoFirmaEstado.objects.get(codigo="SF")
                                firmador.save()
                            SocketMsg(
                                userid=firmador.empleado.persona.usuario.pk,
                                funcpost='refrescarTableros("dbFirmaVB", true)'
                            )
                    context["refrescaTablero"] = True
        return self.render_to_response(context)


class DocumentoEmitirEnviar(TemplateValidaLogin, VistaEdicion):
    template_name = "tramite/documento/emitir/emitir_estado.html"
    model = Documento
    form_class = DocumentoEstadoForm

    def form_valid(self, form):
        context = self.get_context_data(**self.get_form_kwargs())
        if form.is_valid():
            doc = self.object
            if doc.documentoplantilla.documentopdf_set.filter(pdffirma__isnull=False).count() == 0:
                context["error"] = "No encuentra el documento PDF firmado. " \
                                   "Por favor, regréselo a proyecto y fírmelo nuevamente"
            else:
                context["emok"] = True
                codest = form.cleaned_data["codest"]
                if doc.ultimoestado.estado != codest:
                    EmitirDocumento(doc, self.request)
            context["form"] = form
        return self.render_to_response(context)


class DocumentoEmitirObservar(TemplateValidaLogin, VistaEdicion):
    template_name = "tramite/documento/emitir/emitir_observar.html"
    model = Documento
    form_class = DocumentoObservarForm

    def form_valid(self, form):
        context = self.get_context_data(**self.get_form_kwargs())
        if form.is_valid():
            doc = self.get_object()
            DocumentoEstado.objects.create(
                documento=doc,
                estado="OF",
                observacion=form.cleaned_data["observacion"],
                creador=self.request.user
            )
            SocketMsg(
                userid=self.request.user.id,
                funcpost='refrescarTabla("tabladbDespachoO")',
                tipo='primary',
                clase='bg-primary',
                titulo="Correcto",
                mensaje="El documento no se firmó y fue observado correctamente"
            )
            context["observacion_correcta"] = True
        return self.render_to_response(context)


class DocumentoEmitirAnular(TemplateValidaLogin, VistaEdicion):
    template_name = "tramite/documento/emitir/emitir_anular.html"
    model = Documento
    form_class = DocumentoAnularForm

    def form_valid(self, form):
        context = self.get_context_data(**self.get_form_kwargs())
        if form.is_valid():
            doc = self.get_object()
            if doc.numero:
                DocumentoEstado.objects.create(
                    documento=doc,
                    estado="AN",
                    creador=self.request.user,
                    observacion=self.request.POST.get("obs")
                )
                for destino in doc.des_documento.all():
                    if destino.ultimoestado.estado != "AN":
                        DestinoEstado.objects.create(
                            destino=destino,
                            estado="AN",
                            creador=self.request.user
                        )
                context["ok"] = "Su documento se ha anulado correctamente."
            else:
                # Quitamos referencias
                exp = doc.expediente
                doc.expediente = None
                doc.ultimoestado = None
                doc.save()
                # Eliminamos las Plantillas
                for destino in doc.des_documento.all():
                    if hasattr(destino, "documentopdf"):
                        destino.documentopdf.revisiones.all().delete()
                        destino.documentopdf.tokens.all().delete()
                        destino.documentopdf.delete()
                if hasattr(doc, "documentoplantilla"):
                    doc.documentoplantilla.delete()
                # Eliminamos a los Firmantes
                doc.firmas.all().delete()
                destref = []
                # Eliminamos las Referencias
                for ref in doc.referencias.all():
                    if ref.destino:
                        destref.append(ref.destino.id)
                doc.referencias.all().delete()
                for destino in Destino.objects.filter(id__in=destref):
                    if destino.destinoreferencias.count() == 0:
                        ultest = destino.ultimoestado
                        ultest.observacion = "Se anuló la atención: %s." % doc.obtenerNumeroSiglas()
                        ultest.fecha = datetime.now().date()
                        ultest.save()
                        destino.ultimoestado = destino.destinoestados.filter(estado="RE").order_by("-creado").first()
                        destino.save()
                # Eliminamos los Anexos
                for anexo in doc.anexos.all():
                    anexo.firmadores.all().delete()
                    anexo.revisiones.all().delete()
                doc.anexos.all().delete()
                # Eliminamos los Destinos
                doc.des_documento.all().update(ultimoestado=None, ultimoestadomensajeria=None)
                DestinoEstado.objects.filter(destino__in=doc.des_documento.all()).delete()
                DestinoEstadoMensajeria.objects.filter(
                    destino__documento=doc
                ).delete()
                doc.des_documento.all().delete()
                # Eliminamos el Expediente
                if exp and exp.documentos.count() == 0:
                    exp.delete()
                # Eliminamos el documento
                doc.ultimoestado = None
                doc.estadoemitido = None
                doc.save()
                doc.documentoestado_set.all().delete()
                doc.delete()
                context["ok"] = "Su documento se ha eliminado correctamente."
            SocketMsg(
                userid=self.request.user.pk,
                funcpost='refrescarTableros("dbEnProyecto,dbRecepcionados")'
            )
        return self.render_to_response(context)


class DocumentoEmitirAnularEmision(TemplateValidaLogin, VistaEdicion):
    template_name = "tramite/documento/emitir/emitir_anular.html"
    model = Documento
    form_class = DocumentoAnularForm

    def form_valid(self, form):
        context = self.get_context_data(**self.get_form_kwargs())
        if form.is_valid():
            doc = self.get_object()
            if doc.ultimoestado.estado == "EM":
                DocumentoEstado.objects.create(
                    documento=doc,
                    estado="PY",
                    creador=self.request.user
                )
                # Referencias de Respuesta
                for referencia in doc.referencias.filter(destino__isnull=False):
                    if referencia.destino.ultimoestado.estado == "AT":
                        if referencia.destino.destinoreferencias.filter(
                                documento__estadoemitido__isnull=False
                        ).count() == 0:
                            DestinoEstado.objects.create(
                                destino=referencia.destino,
                                estado="RE",
                                creador=self.request.user
                            )
                # Validamos el estado hacia mensajería
                for destino in doc.des_documento.exclude(ultimoestado__estado="AN"):
                    if destino.ultimoestadomensajeria:
                        if destino.ultimoestadomensajeria.estado != "PE":
                            DestinoEstadoMensajeria.objects.create(
                                destino=destino,
                                estado="PE",
                                creador=self.request.user
                            )
                # Notificamos a los destinos de unidades organizacionales
                for destino in doc.des_documento.exclude(ultimoestado__estado="AN"):
                    DestinoEstado.objects.create(
                        destino=destino,
                        estado="RG",
                        creador=self.request.user
                    )
                    if hasattr(destino, "documentopdf"):
                        docpdf = destino.documentopdf
                        docpdf.pdffirma = None
                        docpdf.estado = "G"
                        docpdf.save()
                    if destino.tipodestinatario == "UO":
                        if destino.periodotrabajo.persona.usuario:
                            SocketMsg(
                                userid=destino.periodotrabajo.persona.usuario.pk,
                                funcpost='refrescarTableros("dbEntrada", true)'
                            )
                    elif destino.tipodestinatario in ["CI", "PJ"]:
                        DestinoEstadoMensajeria.objects.create(
                            estado="PE",
                            creador=self.request.user,
                            destino=destino
                        )
                # Notificamos al usuario y al creador
                SocketMsg(
                    userid=self.request.user.pk,
                    funcpost='refrescarTableros("dbEnProyecto,dbEmitidos,dbMiMensajeria", true)'
                )
                if doc.creador != self.request.user:
                    SocketMsg(
                        userid=doc.creador.pk,
                        funcpost='refrescarTableros("dbEnProyecto,dbEmitidos,dbMiMensajeria", true)'
                    )
                context["ok"] = "Se ha anulado la Emisión de su documento correctamente."
            else:
                context["error"] = "No se puede anular la emisión porque ha sido recepcionado por: %s"
                lista = "<ul class='text-left mt-2 mb-0'>"
                for destino in doc.des_documento.filter(ultimoestado__estado__in=["RE", "AT", "AR"]):
                    lista += "<li>%s</li>" % destino.obtenerNombreDestinoPersona()
                lista += "</ul>"
                context["error"] = context["error"] % lista
        return self.render_to_response(context)


class DocumentoEmitirCambiarResponsable(TemplateValidaLogin, VistaEdicion):
    template_name = "tramite/documento/emitir/cambiarresponsable.html"
    model = Documento
    form_class = DocumentoEmitirCambiarResponsableForm
    extra_context = {
        "botonguardartexto": "Cambiar"
    }

    def get_form(self, form_class=None):
        form = super(DocumentoEmitirCambiarResponsable, self).get_form(form_class)
        form.fields["nuevoresponsable"].queryset = TrabajadoresActuales().filter(
            Q(esjefe=True)
            |
            Q(tipo__in=["EN", "EP"])
            |
            Q(area__paracomisiones=True, tipo="NN", cargo__esprincipal=True)
        ).filter(
            area=self.object.responsable.area
        ).order_by("persona__apellidocompleto")
        return form

    def form_valid(self, form):
        context = self.get_context_data(**self.get_form_kwargs())
        if form.is_valid():
            documento = self.get_object()
            documento.responsable = form.cleaned_data["nuevoresponsable"]
            documento.save()
            context["cambio_correcto"] = "El responsable para la firma del documento se ha cambiado correctamente."
        return self.render_to_response(context)
